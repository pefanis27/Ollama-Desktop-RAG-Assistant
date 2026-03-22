from __future__ import annotations

"""
Ollama Desktop RAG Assistant
================================

Πλήρης desktop εφαρμογή Python με GUI για:
- Σύνδεση με τοπικό Ollama
- Έλεγχο αν το Ollama app είναι ανοιχτό
- Εκκίνηση του Ollama app όταν δεν είναι ανοιχτό
- Εκκίνηση του Ollama local API service μέσω ollama CLI
- Αυτόματο διαχωρισμό LLM / Embedding models
- Auto-detect + fallback για επιλογή μοντέλων
- Δημιουργία τοπικών συλλογών γνώσης (RAG)
- Συνομιλία με RAG context
- Επισύναψη αρχείων μέσα στην ίδια την ερώτηση
- Έλεγχο / διόρθωση κώδικα από επισυναπτόμενα αρχεία
- Αποθήκευση corrected code από την απάντηση

Σημείωση:
Η εφαρμογή υλοποιεί RAG και local knowledge augmentation. Δεν κάνει fine-tuning των βαρών
του μοντέλου.
"""

import csv
import io
import json
import os
import re
import shutil
import subprocess
import sys
import time
from dataclasses import asdict, dataclass
from datetime import datetime
import traceback
from pathlib import Path
from typing import Any, Callable, Iterable
from urllib.parse import urljoin

import numpy as np
import psutil
import requests
from docx import Document
from pypdf import PdfReader
from PySide6.QtCore import QObject, QRunnable, Qt, QThreadPool, QTimer, Signal, Slot
from PySide6.QtGui import QAction, QColor, QFont, QTextCursor, QPalette, QTextOption, QCloseEvent
from PySide6.QtWidgets import (
    QApplication,
    QCheckBox,
    QComboBox,
    QDoubleSpinBox,
    QFileDialog,
    QFormLayout,
    QFrame,
    QGridLayout,
    QGroupBox,
    QHBoxLayout,
    QLabel,
    QLineEdit,
    QListWidget,
    QListWidgetItem,
    QMainWindow,
    QMessageBox,
    QPushButton,
    QPlainTextEdit,
    QProgressBar,
    QScrollArea,
    QSizePolicy,
    QSpinBox,
    QSplitter,
    QStatusBar,
    QTabWidget,
    QTextEdit,
    QToolButton,
    QVBoxLayout,
    QWidget,
)


# ============================================================
# Σταθερές εφαρμογής
# ============================================================
APP_NAME = "Ollama Desktop RAG Assistant"
APP_VERSION = "1.2"
DEFAULT_BASE_URL = "http://localhost:11434"
DEFAULT_API_KEY = ""
DEFAULT_TIMEOUT = 180
STATUS_CHECK_TIMEOUT = 1.5
DEFAULT_CHUNK_SIZE = 900
DEFAULT_CHUNK_OVERLAP = 150
DEFAULT_EMBED_BATCH = 16
DEFAULT_TOP_K = 4
DEFAULT_TEMPERATURE = 0.2
DEFAULT_MAX_TOKENS = 8192
DEFAULT_ENABLE_THINKING = False
AUTO_OPTION = "🤖 Αυτόματη επιλογή"
SETTINGS_FILE = Path(__file__).resolve().parent / "app_settings.json"
DATA_DIR = Path(__file__).resolve().parent / "data"
COLLECTIONS_DIR = DATA_DIR / "collections"
COLLECTIONS_DIR.mkdir(parents=True, exist_ok=True)
LOGS_DIR = DATA_DIR / "logs"
LOGS_DIR.mkdir(parents=True, exist_ok=True)
MAX_ATTACHMENT_CHARS = 20000
SUPPORTED_EXTENSIONS = {
    ".txt",
    ".md",
    ".py",
    ".json",
    ".csv",
    ".html",
    ".htm",
    ".pdf",
    ".docx",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".css",
    ".java",
    ".cpp",
    ".c",
    ".cs",
    ".go",
    ".rs",
    ".sql",
    ".yaml",
    ".yml",
    ".xml",
    ".ini",
    ".toml",
    ".log",
}

EMBEDDING_NAME_HINTS = (
    "embedding",
    "embed",
    "embeddinggemma",
    "qwen3-embedding",
    "all-minilm",
    "bge",
    "e5",
    "gte",
    "mxbai",
    "nomic-embed",
    "snowflake-arctic-embed",
    "jina-embeddings",
    "minilm",
    "gte-",
)

CODE_FILE_EXTENSIONS = {
    ".py",
    ".js",
    ".ts",
    ".tsx",
    ".jsx",
    ".java",
    ".cpp",
    ".c",
    ".cs",
    ".go",
    ".rs",
    ".php",
    ".rb",
    ".swift",
    ".kt",
    ".sql",
    ".html",
    ".css",
    ".xml",
    ".yaml",
    ".yml",
    ".json",
}

CHAT_MODE_RAG = "RAG με βάση γνώσης"
CHAT_MODE_PLAIN = "Συνομιλία χωρίς RAG"
MODEL_LIST_FILTER_MODE = "mode"
MODEL_LIST_FILTER_LLM = "llm"
MODEL_LIST_FILTER_EMBEDDING = "embedding"
MODEL_LIST_FILTER_ALL = "all"

MODEL_LOAD_SCOPE_ACTIVE = "active"
MODEL_LOAD_SCOPE_LLM = "llm"
MODEL_LOAD_SCOPE_EMBEDDING = "embedding"
MODEL_LOAD_SCOPE_BOTH = "both"


# ============================================================
# Exceptions
# ============================================================
class AppError(Exception):
    pass


class DocumentLoaderError(AppError):
    pass


class OllamaError(AppError):
    pass


class RAGEngineError(AppError):
    pass


# ============================================================
# Dataclasses
# ============================================================
@dataclass(slots=True)
class LoadedDocument:
    filename: str
    text: str
    source_path: str | None = None


@dataclass(slots=True)
class TextChunk:
    chunk_id: str
    filename: str
    text: str
    chunk_index: int


@dataclass(slots=True)
class SearchResult:
    score: float
    filename: str
    chunk_id: str
    chunk_index: int
    text: str


@dataclass(slots=True)
class AvailableModel:
    identifier: str
    type: str
    display_name: str | None = None
    publisher: str | None = None
    arch: str | None = None
    is_loaded: bool = False

    @property
    def label(self) -> str:
        title = self.display_name or self.identifier
        suffix = " (loaded)" if self.is_loaded else ""
        return f"{title} [{self.type}]{suffix}"


@dataclass(slots=True)
class AppSettings:
    base_url: str = DEFAULT_BASE_URL
    api_key: str = DEFAULT_API_KEY
    timeout_seconds: int = DEFAULT_TIMEOUT
    ollama_exe_path: str = ""
    ollama_cli_path: str = ""
    auto_open_ollama_on_start: bool = False
    auto_start_service_on_start: bool = False
    auto_refresh_models_on_start: bool = False
    auto_load_selected_models: bool = False
    model_load_scope: str = MODEL_LOAD_SCOPE_ACTIVE
    preferred_llm: str = AUTO_OPTION
    preferred_embedding: str = AUTO_OPTION
    last_collection_name: str = "knowledge_base"
    chunk_size: int = DEFAULT_CHUNK_SIZE
    chunk_overlap: int = DEFAULT_CHUNK_OVERLAP
    embed_batch_size: int = DEFAULT_EMBED_BATCH
    top_k: int = DEFAULT_TOP_K
    temperature: float = DEFAULT_TEMPERATURE
    max_tokens: int = DEFAULT_MAX_TOKENS
    last_chat_mode: str = CHAT_MODE_PLAIN


# ============================================================
# Ρυθμίσεις / persistence
# ============================================================
def load_settings() -> AppSettings:
    if not SETTINGS_FILE.exists():
        return AppSettings()
    try:
        payload = json.loads(SETTINGS_FILE.read_text(encoding="utf-8"))
        migrations = {
            "lmstudio_exe_path": "ollama_exe_path",
            "lmstudio_cli_path": "ollama_cli_path",
            "auto_open_lmstudio_on_start": "auto_open_ollama_on_start",
            "auto_start_service_on_start": "auto_start_service_on_start",
        }
        for old_key, new_key in migrations.items():
            if old_key in payload and new_key not in payload:
                payload[new_key] = payload.pop(old_key)
        if payload.get("base_url") == "http://localhost:1234":
            payload["base_url"] = DEFAULT_BASE_URL
        if payload.get("api_key") == "lm-studio":
            payload["api_key"] = DEFAULT_API_KEY
        allowed_fields = set(AppSettings.__dataclass_fields__.keys())
        payload = {key: value for key, value in payload.items() if key in allowed_fields}
        return AppSettings(**payload)
    except Exception:
        return AppSettings()


def save_settings(settings: AppSettings) -> None:
    SETTINGS_FILE.write_text(
        json.dumps(asdict(settings), ensure_ascii=False, indent=2),
        encoding="utf-8",
    )


# ============================================================
# Βοηθητικές συναρτήσεις εγγράφων
# ============================================================
def normalize_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"\r\n?", "\n", text)
    text = re.sub(r"[\t\f\v]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    text = re.sub(r" {2,}", " ", text)
    return text.strip()


def is_supported_file(filename: str) -> bool:
    return Path(filename).suffix.lower() in SUPPORTED_EXTENSIONS


def read_pdf(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages: list[str] = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages)


def read_docx(file_bytes: bytes) -> str:
    doc = Document(io.BytesIO(file_bytes))
    return "\n".join(paragraph.text for paragraph in doc.paragraphs)


def read_csv(file_bytes: bytes) -> str:
    text = file_bytes.decode("utf-8", errors="ignore")
    rows = csv.reader(io.StringIO(text))
    return "\n".join(" | ".join(cell.strip() for cell in row) for row in rows)


def read_json(file_bytes: bytes) -> str:
    obj = json.loads(file_bytes.decode("utf-8", errors="ignore"))
    return json.dumps(obj, ensure_ascii=False, indent=2)


def read_text(file_bytes: bytes) -> str:
    return file_bytes.decode("utf-8", errors="ignore")


def extract_text_from_bytes(filename: str, file_bytes: bytes) -> LoadedDocument:
    suffix = Path(filename).suffix.lower()
    if suffix not in SUPPORTED_EXTENSIONS:
        raise DocumentLoaderError(f"Μη υποστηριζόμενος τύπος αρχείου: {suffix}")

    try:
        if suffix == ".pdf":
            text = read_pdf(file_bytes)
        elif suffix == ".docx":
            text = read_docx(file_bytes)
        elif suffix == ".csv":
            text = read_csv(file_bytes)
        elif suffix == ".json":
            text = read_json(file_bytes)
        else:
            text = read_text(file_bytes)
    except Exception as exc:
        raise DocumentLoaderError(f"Αποτυχία ανάγνωσης του αρχείου '{filename}': {exc}") from exc

    text = normalize_text(text)
    if not text:
        raise DocumentLoaderError(f"Το αρχείο '{filename}' δεν περιέχει αναγνώσιμο κείμενο.")

    return LoadedDocument(filename=filename, text=text)


def extract_text_from_path(path: str | Path) -> LoadedDocument:
    path = Path(path)
    file_bytes = path.read_bytes()
    loaded = extract_text_from_bytes(path.name, file_bytes)
    return LoadedDocument(filename=loaded.filename, text=loaded.text, source_path=str(path))


def split_text(text: str, chunk_size: int, chunk_overlap: int) -> list[str]:
    if chunk_size <= 0:
        raise ValueError("Το chunk_size πρέπει να είναι θετικός αριθμός.")
    if chunk_overlap < 0:
        raise ValueError("Το chunk_overlap δεν μπορεί να είναι αρνητικό.")
    if chunk_overlap >= chunk_size:
        raise ValueError("Το chunk_overlap πρέπει να είναι μικρότερο από το chunk_size.")

    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    text_len = len(text)

    while start < text_len:
        end = min(start + chunk_size, text_len)
        candidate = text[start:end]

        if end < text_len:
            breakpoints = [
                candidate.rfind("\n\n"),
                candidate.rfind("\n"),
                candidate.rfind(". "),
                candidate.rfind(" "),
            ]
            best_break = max(breakpoints)
            if best_break > chunk_size * 0.55:
                end = start + best_break + 1
                candidate = text[start:end]

        chunks.append(candidate.strip())
        if end >= text_len:
            break
        start = max(0, end - chunk_overlap)

    return [chunk for chunk in chunks if chunk]


def build_chunks(documents: Iterable[LoadedDocument], chunk_size: int, chunk_overlap: int) -> list[TextChunk]:
    chunks: list[TextChunk] = []
    for document in documents:
        parts = split_text(document.text, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        for idx, part in enumerate(parts):
            chunks.append(
                TextChunk(
                    chunk_id=f"{Path(document.filename).stem}_{idx:04d}",
                    filename=document.filename,
                    text=part,
                    chunk_index=idx,
                )
            )
    return chunks


def truncate_middle(text: str, max_chars: int) -> str:
    if len(text) <= max_chars:
        return text
    half = max_chars // 2
    return text[:half] + "\n\n...[ΠΕΡΙΕΧΟΜΕΝΟ ΠΕΡΙΚΟΠΗΚΕ]...\n\n" + text[-half:]


def truncate_ui_name(text: str, max_chars: int = 54) -> str:
    text = (text or "").strip()
    if len(text) <= max_chars:
        return text
    return text[: max_chars - 3].rstrip() + "..."


def infer_code_language(filename: str) -> str:
    suffix = Path(filename).suffix.lower().lstrip(".")
    if not suffix:
        return "text"
    return suffix


def extract_first_code_block(text: str) -> tuple[str | None, str | None]:
    match = re.search(r"```([\w#+\-.]*)\n(.*?)```", text, flags=re.DOTALL)
    if not match:
        return None, None
    language = match.group(1).strip() or None
    code = match.group(2).rstrip()
    return language, code


# ============================================================
# Ollama επικοινωνία
# ============================================================
class OllamaClient:
    def __init__(self, base_url: str, api_key: str = DEFAULT_API_KEY, timeout: int = DEFAULT_TIMEOUT) -> None:
        self.base_url = base_url.rstrip("/")
        self.api_key = api_key
        self.timeout = timeout
        self.session = requests.Session()

    @property
    def headers(self) -> dict[str, str]:
        headers = {"Content-Type": "application/json"}
        if self.api_key:
            headers["Authorization"] = f"Bearer {self.api_key}"
        return headers

    def _classify_model_type(self, identifier: str, capabilities: list[str] | None) -> str:
        caps = {str(x).strip().lower() for x in (capabilities or []) if str(x).strip()}
        searchable = identifier.lower()
        if "embedding" in caps and "completion" not in caps:
            return "embedding"
        if "completion" in caps or "vision" in caps or "tools" in caps:
            return "llm"
        if any(hint in searchable for hint in EMBEDDING_NAME_HINTS):
            return "embedding"
        return "llm"

    @staticmethod
    def _preview_value(value: Any, limit: int = 700) -> str:
        try:
            if isinstance(value, (dict, list)):
                text = json.dumps(value, ensure_ascii=False)
            else:
                text = str(value)
        except Exception:
            text = repr(value)
        text = text.replace("\r", " ").replace("\n", " ").strip()
        return text if len(text) <= limit else text[:limit] + "..."

    @staticmethod
    def _coerce_text(value: Any) -> str:
        if value is None:
            return ""
        if isinstance(value, str):
            return value.strip()
        if isinstance(value, list):
            parts = [OllamaClient._coerce_text(item) for item in value]
            return "\n".join(part for part in parts if part).strip()
        if isinstance(value, dict):
            for key in ("content", "response", "text", "message", "output"):
                coerced = OllamaClient._coerce_text(value.get(key))
                if coerced:
                    return coerced
            return ""
        return str(value).strip()

    @classmethod
    def _describe_payload(cls, payload: Any) -> str:
        if not isinstance(payload, dict):
            return cls._preview_value(payload)
        summary: dict[str, Any] = {}
        for key in ("model", "done", "done_reason", "total_duration", "load_duration", "prompt_eval_count", "eval_count"):
            if key in payload:
                summary[key] = payload.get(key)
        message = payload.get("message")
        if isinstance(message, dict):
            summary["message_keys"] = sorted(message.keys())
            summary["message_content_preview"] = cls._preview_value(message.get("content") or message.get("response") or "")
            summary["message_thinking_preview"] = cls._preview_value(message.get("thinking") or "")
            if message.get("tool_calls"):
                summary["tool_calls"] = message.get("tool_calls")
        for key in ("response", "content", "text", "thinking", "error"):
            if key in payload:
                summary[key] = payload.get(key)
        return cls._preview_value(summary)

    def _parse_json_response(self, response: requests.Response, method: str, path: str) -> Any:
        text = response.text or ""
        if not text.strip():
            return {}
        try:
            return response.json()
        except ValueError as exc:
            raise OllamaError(
                f"Αποτυχία parsing JSON από Ollama ({method} {path}) | status={response.status_code} | body={self._preview_value(text)}"
            ) from exc

    def _get(self, path: str, timeout: float | None = None) -> Any:
        url = urljoin(self.base_url + "/", path.lstrip("/"))
        response = None
        try:
            response = self.session.get(url, headers=self.headers, timeout=timeout or self.timeout)
            response.raise_for_status()
            return self._parse_json_response(response, "GET", path)
        except requests.RequestException as exc:
            details = ""
            if response is not None:
                details = f" | status={response.status_code} | body={self._preview_value(response.text)}"
            raise OllamaError(f"Αποτυχία GET στο Ollama ({path}): {exc}{details}") from exc

    def _post(self, path: str, payload: dict[str, Any], timeout: float | None = None) -> Any:
        url = urljoin(self.base_url + "/", path.lstrip("/"))
        response = None
        try:
            response = self.session.post(url, json=payload, headers=self.headers, timeout=timeout or self.timeout)
            response.raise_for_status()
            return self._parse_json_response(response, "POST", path)
        except requests.RequestException as exc:
            details = f" | payload={self._preview_value(payload)}"
            try:
                if response is not None:
                    details += f" | status={response.status_code} | body={self._preview_value(response.text)}"
            except Exception:
                pass
            raise OllamaError(f"Αποτυχία POST στο Ollama ({path}): {exc}{details}") from exc

    def _post_stream_jsonl(
        self,
        path: str,
        payload: dict[str, Any],
        timeout: float | None = None,
    ) -> list[dict[str, Any]]:
        url = urljoin(self.base_url + "/", path.lstrip("/"))
        response = None
        events: list[dict[str, Any]] = []
        bad_lines: list[str] = []
        try:
            response = self.session.post(
                url,
                json=payload,
                headers=self.headers,
                timeout=timeout or self.timeout,
                stream=True,
            )
            response.raise_for_status()
            for raw_line in response.iter_lines(decode_unicode=True):
                if not raw_line:
                    continue
                line = raw_line.strip()
                if not line:
                    continue
                try:
                    item = json.loads(line)
                except json.JSONDecodeError:
                    if len(bad_lines) < 5:
                        bad_lines.append(line[:200])
                    continue
                if isinstance(item, dict):
                    events.append(item)
            if not events and bad_lines:
                raise OllamaError(
                    f"Το Ollama επέστρεψε stream χωρίς έγκυρα JSON events ({path}) | bad_lines={bad_lines} | payload={self._preview_value(payload)}"
                )
            return events
        except requests.RequestException as exc:
            details = f" | payload={self._preview_value(payload)}"
            try:
                if response is not None:
                    details += f" | status={response.status_code} | body={self._preview_value(response.text)}"
            except Exception:
                pass
            raise OllamaError(f"Αποτυχία POST stream στο Ollama ({path}): {exc}{details}") from exc

    @staticmethod
    def _message_content_from_payload(data: dict[str, Any]) -> str:
        message = data.get("message") or {}
        text = OllamaClient._coerce_text(message)
        if text:
            return text
        return OllamaClient._coerce_text(data)

    @staticmethod
    def _message_thinking_from_payload(data: dict[str, Any]) -> str:
        message = data.get("message") or {}
        if isinstance(message, dict):
            thinking = OllamaClient._coerce_text(message.get("thinking"))
            if thinking:
                return thinking
        return OllamaClient._coerce_text(data.get("thinking"))

    @staticmethod
    def _build_inference_options(temperature: float, max_tokens: int, enable_thinking: bool = DEFAULT_ENABLE_THINKING) -> dict[str, Any]:
        options: dict[str, Any] = {
            "temperature": temperature,
            "num_predict": max_tokens,
        }
        if not enable_thinking:
            options["enable_thinking"] = False
        return options

    def _build_chat_payload(
        self,
        model: str,
        messages: list[dict[str, str]],
        temperature: float,
        max_tokens: int,
        *,
        stream: bool,
        enable_thinking: bool = DEFAULT_ENABLE_THINKING,
    ) -> dict[str, Any]:
        return {
            "model": model,
            "messages": messages,
            "stream": stream,
            "think": bool(enable_thinking),
            "keep_alive": "15m",
            "options": self._build_inference_options(temperature, max_tokens, enable_thinking=enable_thinking),
        }

    def _build_generate_payload(
        self,
        model: str,
        messages: list[dict[str, str]],
        temperature: float,
        max_tokens: int,
        *,
        enable_thinking: bool = DEFAULT_ENABLE_THINKING,
    ) -> dict[str, Any]:
        return {
            "model": model,
            "prompt": self._messages_to_prompt(messages),
            "stream": False,
            "think": bool(enable_thinking),
            "keep_alive": "15m",
            "options": self._build_inference_options(temperature, max_tokens, enable_thinking=enable_thinking),
        }

    @staticmethod
    def _force_final_answer_messages(messages: list[dict[str, str]]) -> list[dict[str, str]]:
        cloned = [dict(item) for item in messages]
        if not cloned:
            return cloned
        last = dict(cloned[-1])
        extra_instruction = (
            "Σημαντικό: Δώσε μόνο την τελική απάντηση χωρίς thinking, reasoning ή εσωτερική ανάλυση. "
            "Αν χρειάζεται κώδικας, δώσε απευθείας τον τελικό κώδικα."
        )
        last["content"] = (extra_instruction + "\n\n" + str(last.get('content') or '').strip()).strip()
        cloned[-1] = last
        return cloned

    @staticmethod
    def _format_tool_calls(tool_calls: list[dict[str, Any]]) -> str:
        lines = ["Το μοντέλο επέστρεψε tool calls αντί για τελικό κείμενο:"]
        for idx, call in enumerate(tool_calls, start=1):
            function = (call or {}).get("function") or {}
            name = str(function.get("name") or f"tool_{idx}").strip()
            arguments = function.get("arguments")
            if isinstance(arguments, (dict, list)):
                args_text = json.dumps(arguments, ensure_ascii=False, indent=2)
            else:
                args_text = str(arguments or "").strip()
            lines.append(f"{idx}. {name}")
            if args_text:
                lines.append(args_text)
        return "\n".join(lines).strip()

    @staticmethod
    def _messages_to_prompt(messages: list[dict[str, str]]) -> str:
        chunks: list[str] = []
        for item in messages:
            role = str(item.get("role") or "user").strip().upper()
            content = str(item.get("content") or "").strip()
            if content:
                chunks.append(f"[{role}]\n{content}")
        chunks.append("[ASSISTANT]\n")
        return "\n\n".join(chunks).strip()

    def _chat_streaming_content(
        self,
        model: str,
        messages: list[dict[str, str]],
        temperature: float,
        max_tokens: int,
    ) -> str:
        payload = self._build_chat_payload(
            model,
            messages,
            temperature,
            max_tokens,
            stream=True,
            enable_thinking=DEFAULT_ENABLE_THINKING,
        )
        events = self._post_stream_jsonl("/api/chat", payload, timeout=max(180, self.timeout))
        pieces: list[str] = []
        last_tool_calls: list[dict[str, Any]] = []
        for event in events:
            message = event.get("message") or {}
            if not isinstance(message, dict):
                message = {}
            chunk = str(message.get("content") or "")
            if chunk:
                pieces.append(chunk)
            tool_calls = message.get("tool_calls") or []
            if tool_calls:
                last_tool_calls = tool_calls
        content = "".join(pieces).strip()
        if content:
            return content
        if last_tool_calls:
            return self._format_tool_calls(last_tool_calls)
        return ""

    def _generate_fallback_content(
        self,
        model: str,
        messages: list[dict[str, str]],
        temperature: float,
        max_tokens: int,
    ) -> str:
        payload = self._build_generate_payload(
            model,
            messages,
            temperature,
            max_tokens,
            enable_thinking=DEFAULT_ENABLE_THINKING,
        )
        data = self._post("/api/generate", payload, timeout=max(180, self.timeout))
        return self._coerce_text(data)

    def is_server_running(self, timeout: float = STATUS_CHECK_TIMEOUT) -> bool:
        try:
            self._get("/api/version", timeout=timeout)
            return True
        except OllamaError:
            return False

    def get_version(self) -> str:
        data = self._get("/api/version")
        return str(data.get("version") or "").strip()

    def list_running_models(self) -> set[str]:
        try:
            data = self._get("/api/ps", timeout=max(2.0, STATUS_CHECK_TIMEOUT))
        except OllamaError:
            return set()
        running: set[str] = set()
        for item in data.get("models", []) or []:
            name = str(item.get("name") or item.get("model") or "").strip()
            if name:
                running.add(name)
        return running

    def show_model_details(self, model_name: str) -> dict[str, Any]:
        return self._post("/api/show", {"model": model_name}, timeout=max(10, self.timeout))

    def list_models(self, progress: Callable[[str], None] | None = None) -> list[AvailableModel]:
        progress = progress or (lambda _msg: None)
        payload = self._get("/api/tags", timeout=max(10, self.timeout))
        models = payload.get("models", []) or []
        running = self.list_running_models()
        parsed: list[AvailableModel] = []
        total = len(models)

        for idx, item in enumerate(models, start=1):
            identifier = str(item.get("name") or item.get("model") or "").strip()
            if not identifier:
                continue

            capabilities: list[str] = []
            try:
                details_payload = self.show_model_details(identifier)
                capabilities = [str(x) for x in (details_payload.get("capabilities") or [])]
            except Exception:
                capabilities = []

            model_type = self._classify_model_type(identifier, capabilities)
            details = item.get("details") or {}
            families = details.get("families") or []
            arch = details.get("family") or (families[0] if families else None)

            if total <= 12 or idx in {1, total} or idx % max(1, total // 6) == 0:
                progress(f"🔎 Έλεγχος model {idx}/{total}: {identifier}")

            parsed.append(
                AvailableModel(
                    identifier=identifier,
                    type=model_type,
                    display_name=identifier,
                    publisher=None,
                    arch=arch,
                    is_loaded=identifier in running,
                )
            )

        parsed.sort(key=lambda m: (m.type, 0 if m.is_loaded else 1, (m.display_name or m.identifier).lower()))
        return parsed

    def load_model(self, model_identifier: str, model_type: str) -> dict[str, Any]:
        keep_alive = "15m"
        if model_type == "embedding":
            return self._post(
                "/api/embed",
                {
                    "model": model_identifier,
                    "input": "warmup",
                    "truncate": True,
                    "keep_alive": keep_alive,
                },
                timeout=max(60, self.timeout),
            )
        return self._post(
            "/api/chat",
            {
                "model": model_identifier,
                "messages": [{"role": "user", "content": "OK"}],
                "stream": False,
                "think": False,
                "keep_alive": keep_alive,
                "options": {"temperature": 0.0, "num_predict": 1, "enable_thinking": False},
            },
            timeout=max(60, self.timeout),
        )

    def create_embeddings(self, texts: list[str], model: str) -> list[list[float]]:
        if not texts:
            return []
        payload = {
            "model": model,
            "input": [text.replace("\n", " ") for text in texts],
            "truncate": True,
            "keep_alive": "15m",
        }
        data = self._post("/api/embed", payload, timeout=max(120, self.timeout))
        vectors = data.get("embeddings", []) or []
        if not vectors:
            raise OllamaError("Το Ollama δεν επέστρεψε embeddings.")
        return vectors

    def chat(
        self,
        model: str,
        messages: list[dict[str, str]],
        temperature: float = DEFAULT_TEMPERATURE,
        max_tokens: int = DEFAULT_MAX_TOKENS,
    ) -> str:
        payload = self._build_chat_payload(
            model,
            messages,
            temperature,
            max_tokens,
            stream=False,
            enable_thinking=DEFAULT_ENABLE_THINKING,
        )
        errors: list[str] = []
        retry_messages: list[dict[str, str]] | None = None

        try:
            data = self._post("/api/chat", payload, timeout=max(180, self.timeout))
            content = self._message_content_from_payload(data)
            if content:
                return content
            tool_calls = ((data.get("message") or {}) if isinstance(data, dict) else {}).get("tool_calls") or []
            if tool_calls:
                return self._format_tool_calls(tool_calls)

            thinking = self._message_thinking_from_payload(data)
            if thinking:
                retry_messages = self._force_final_answer_messages(messages)
                errors.append(
                    "thinking χωρίς τελικό content από /api/chat | " + self._describe_payload(data)
                )
            else:
                errors.append("κενή non-stream απόκριση από /api/chat | " + self._describe_payload(data))
        except OllamaError as exc:
            errors.append(str(exc))

        chat_messages = retry_messages or messages

        try:
            streamed_content = self._chat_streaming_content(model, chat_messages, temperature, max_tokens)
            if streamed_content:
                return streamed_content
            errors.append("κενή streamed απόκριση από /api/chat | δεν βρέθηκε κείμενο στα streamed events")
        except OllamaError as exc:
            errors.append(str(exc))

        try:
            generated_content = self._generate_fallback_content(model, chat_messages, temperature, max_tokens)
            if generated_content:
                return generated_content
            errors.append("κενή fallback απόκριση από /api/generate | δεν βρέθηκε πεδίο response/content")
        except OllamaError as exc:
            errors.append(str(exc))

        details = " | ".join(err for err in errors if err)
        if details:
            raise OllamaError(f"Το Ollama δεν επέστρεψε αξιοποιήσιμη απάντηση. {details}")
        raise OllamaError("Το Ollama δεν επέστρεψε αξιοποιήσιμη απάντηση.")


# ============================================================
# Controller για Ollama app / service
# ============================================================
class OllamaController:
    def __init__(self, settings: AppSettings) -> None:
        self.settings = settings

    def client(self) -> OllamaClient:
        return OllamaClient(
            base_url=self.settings.base_url,
            api_key=self.settings.api_key,
            timeout=self.settings.timeout_seconds,
        )

    def discover_executable_path(self) -> str:
        candidates: list[Path] = []
        home = Path.home()

        if sys.platform.startswith("win"):
            env_paths = [
                os.environ.get("LOCALAPPDATA"),
                os.environ.get("ProgramFiles"),
                os.environ.get("ProgramFiles(x86)"),
            ]
            for root in env_paths:
                if root:
                    candidates.extend([
                        Path(root) / "Programs" / "Ollama" / "ollama app.exe",
                        Path(root) / "Programs" / "Ollama" / "Ollama app.exe",
                        Path(root) / "Programs" / "Ollama" / "ollama.exe",
                        Path(root) / "Programs" / "Ollama" / "Ollama.exe",
                        Path(root) / "Ollama" / "ollama app.exe",
                        Path(root) / "Ollama" / "ollama.exe",
                    ])
        elif sys.platform == "darwin":
            candidates.append(Path("/Applications/Ollama.app/Contents/MacOS/Ollama"))
        else:
            found = shutil.which("ollama")
            if found:
                candidates.append(Path(found))
            candidates.append(home / ".local" / "bin" / "ollama")
            candidates.append(Path("/usr/bin/ollama"))

        for candidate in candidates:
            if candidate.exists():
                return str(candidate)
        return ""

    def discover_cli_path(self) -> str:
        if sys.platform.startswith("win"):
            candidates = [shutil.which("ollama"), shutil.which("ollama.exe"), shutil.which("ollama.cmd")]
        else:
            candidates = [shutil.which("ollama")]
        for candidate in candidates:
            if candidate:
                return candidate
        return ""

    def get_effective_exe_path(self) -> str:
        configured = self.settings.ollama_exe_path.strip()
        if configured and Path(configured).exists():
            return configured
        discovered = self.discover_executable_path()
        if discovered:
            self.settings.ollama_exe_path = discovered
        return discovered

    def get_effective_cli_path(self) -> str:
        configured = self.settings.ollama_cli_path.strip()
        if configured and (Path(configured).exists() or shutil.which(configured)):
            return configured
        discovered = self.discover_cli_path()
        if discovered:
            self.settings.ollama_cli_path = discovered
        return discovered

    def is_ollama_app_running(self) -> bool:
        keywords = ("ollama app", "ollama")
        target_path = self.get_effective_exe_path().lower()

        for proc in psutil.process_iter(["name", "exe", "cmdline"]):
            try:
                name = (proc.info.get("name") or "").lower()
                exe = (proc.info.get("exe") or "").lower()
                cmdline = " ".join(proc.info.get("cmdline") or []).lower()
                if any(keyword in name for keyword in keywords):
                    return True
                if any(keyword in cmdline for keyword in keywords):
                    return True
                if target_path and exe == target_path:
                    return True
            except (psutil.NoSuchProcess, psutil.AccessDenied):
                continue
        return False

    def open_ollama_app(self) -> str:
        exe_path = self.get_effective_exe_path()
        if not exe_path:
            raise OllamaError("Δεν βρέθηκε το εκτελέσιμο του Ollama. Ρύθμισε το path από το GUI.")
        try:
            if sys.platform.startswith("win"):
                os.startfile(exe_path)  # type: ignore[attr-defined]
            else:
                subprocess.Popen([exe_path])
            return exe_path
        except Exception as exc:
            raise OllamaError(f"Αποτυχία εκκίνησης του Ollama: {exc}") from exc

    def _detached_popen_kwargs(self) -> dict[str, Any]:
        if sys.platform.startswith("win"):
            creationflags = 0
            creationflags |= getattr(subprocess, "DETACHED_PROCESS", 0)
            creationflags |= getattr(subprocess, "CREATE_NEW_PROCESS_GROUP", 0)
            return {"creationflags": creationflags}
        return {"start_new_session": True}

    def start_service(self) -> str:
        cli_path = self.get_effective_cli_path()
        if not cli_path:
            raise OllamaError("Δεν βρέθηκε το ollama CLI. Εγκατέστησέ το ή δήλωσε το path του από το GUI.")
        cmd = [cli_path, "serve"]
        try:
            subprocess.Popen(
                cmd,
                cwd=str(Path(__file__).resolve().parent),
                stdout=subprocess.DEVNULL,
                stderr=subprocess.DEVNULL,
                **self._detached_popen_kwargs(),
            )
            return " ".join(cmd)
        except Exception as exc:
            raise OllamaError(f"Αποτυχία εκκίνησης του Ollama service: {exc}") from exc

    def wait_for_server(self, timeout_seconds: int = 20) -> bool:
        client = self.client()
        started = time.time()
        while time.time() - started < timeout_seconds:
            if client.is_server_running():
                return True
            time.sleep(1.0)
        return False

    def ensure_ready(self, progress: Callable[[str], None] | None = None) -> dict[str, bool]:
        progress = progress or (lambda _msg: None)
        app_running_before = self.is_ollama_app_running()
        server_running_before = self.client().is_server_running()
        app_opened = False
        server_started = False

        if not app_running_before and self.settings.auto_open_ollama_on_start:
            progress("🔎 Το Ollama δεν φαίνεται ανοιχτό. Προσπάθεια εκκίνησης...")
            self.open_ollama_app()
            app_opened = True
            time.sleep(3)

        if not server_running_before and self.settings.auto_start_service_on_start:
            progress("🌐 Το service του Ollama δεν απαντά. Προσπάθεια εκκίνησης μέσω CLI...")
            self.start_service()
            server_started = self.wait_for_server(timeout_seconds=20)

        return {
            "app_running": self.is_ollama_app_running(),
            "server_running": self.client().is_server_running(),
            "app_opened": app_opened,
            "server_started": server_started,
        }

    def ensure_server_running(self, progress: Callable[[str], None] | None = None, timeout_seconds: int = 20) -> bool:
        progress = progress or (lambda _msg: None)
        client = self.client()
        if client.is_server_running():
            return True

        if self.settings.auto_open_ollama_on_start and not self.is_ollama_app_running():
            progress("🚀 Το Ollama app δεν είναι ανοιχτό. Προσπάθεια εκκίνησης...")
            self.open_ollama_app()
            time.sleep(2.5)
            if client.is_server_running():
                return True

        if self.settings.auto_start_service_on_start:
            progress("🌐 Το Ollama service δεν απαντά. Προσπάθεια εκκίνησης...")
            self.start_service()
            if self.wait_for_server(timeout_seconds=timeout_seconds):
                return True

        return client.is_server_running(timeout=STATUS_CHECK_TIMEOUT)


# ============================================================
# RAG Engine
# ============================================================
class RAGEngine:
    def __init__(self, collections_dir: str | Path = COLLECTIONS_DIR) -> None:
        self.collections_dir = Path(collections_dir)
        self.collections_dir.mkdir(parents=True, exist_ok=True)

    def list_collections(self) -> list[str]:
        return sorted(item.name for item in self.collections_dir.iterdir() if item.is_dir())

    def collection_dir(self, collection_name: str) -> Path:
        safe_name = collection_name.strip().replace(" ", "_")
        if not safe_name:
            raise RAGEngineError("Το όνομα της συλλογής δεν μπορεί να είναι κενό.")
        return self.collections_dir / safe_name

    def delete_collection(self, collection_name: str) -> None:
        path = self.collection_dir(collection_name)
        if path.exists():
            shutil.rmtree(path)

    def get_collection_metadata(self, collection_name: str) -> dict[str, Any]:
        metadata_path = self.collection_dir(collection_name) / "metadata.json"
        if not metadata_path.exists():
            raise RAGEngineError(f"Δεν βρέθηκε metadata για τη συλλογή '{collection_name}'.")
        return json.loads(metadata_path.read_text(encoding="utf-8"))

    def build_collection(
        self,
        collection_name: str,
        file_paths: list[str],
        client: OllamaClient,
        embedding_model: str,
        chunk_size: int,
        chunk_overlap: int,
        batch_size: int,
        progress: Callable[[str], None] | None = None,
    ) -> dict[str, Any]:
        progress = progress or (lambda _msg: None)
        collection_dir = self.collection_dir(collection_name)
        source_dir = collection_dir / "source_files"
        source_dir.mkdir(parents=True, exist_ok=True)

        documents: list[LoadedDocument] = []
        saved_files: list[str] = []

        for path_str in file_paths:
            path = Path(path_str)
            if not path.exists():
                raise RAGEngineError(f"Το αρχείο δεν βρέθηκε: {path}")
            progress(f"📄 Ανάγνωση αρχείου: {path.name}")
            document = extract_text_from_path(path)
            shutil.copy2(path, source_dir / path.name)
            documents.append(document)
            saved_files.append(path.name)

        if not documents:
            raise RAGEngineError("Δεν επιλέχθηκαν αρχεία για δημιουργία γνώσης.")

        progress("✂️ Δημιουργία chunks...")
        chunks = build_chunks(documents, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        if not chunks:
            raise RAGEngineError("Δεν προέκυψαν chunks από τα επιλεγμένα αρχεία.")

        all_embeddings: list[list[float]] = []
        chunk_texts = [chunk.text for chunk in chunks]
        total_batches = max(1, (len(chunk_texts) + batch_size - 1) // batch_size)

        progress_step = max(1, total_batches // 10)
        for batch_index, start in enumerate(range(0, len(chunk_texts), batch_size), start=1):
            batch = chunk_texts[start : start + batch_size]
            if batch_index == 1 or batch_index == total_batches or batch_index % progress_step == 0:
                progress(f"🧠 Δημιουργία embeddings batch {batch_index}/{total_batches}...")
            vectors = client.create_embeddings(batch, model=embedding_model)
            all_embeddings.extend(vectors)

        matrix = np.array(all_embeddings, dtype=np.float32)
        normalized = self.normalize(matrix)

        metadata = {
            "collection_name": collection_dir.name,
            "created_at": datetime.now().isoformat(timespec="seconds"),
            "embedding_model": embedding_model,
            "chunk_size": chunk_size,
            "chunk_overlap": chunk_overlap,
            "documents": saved_files,
            "num_chunks": len(chunks),
            "embedding_dim": int(normalized.shape[1]),
        }

        chunks_payload = [
            {
                "chunk_id": chunk.chunk_id,
                "filename": chunk.filename,
                "chunk_index": chunk.chunk_index,
                "text": chunk.text,
            }
            for chunk in chunks
        ]

        np.save(collection_dir / "embeddings.npy", normalized)
        (collection_dir / "chunks.json").write_text(
            json.dumps(chunks_payload, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )
        (collection_dir / "metadata.json").write_text(
            json.dumps(metadata, ensure_ascii=False, indent=2),
            encoding="utf-8",
        )

        progress("✅ Η συλλογή δημιουργήθηκε επιτυχώς.")
        return metadata

    def search(
        self,
        collection_name: str,
        query: str,
        client: OllamaClient,
        embedding_model: str,
        top_k: int,
    ) -> list[SearchResult]:
        collection_dir = self.collection_dir(collection_name)
        chunks_path = collection_dir / "chunks.json"
        vectors_path = collection_dir / "embeddings.npy"

        if not chunks_path.exists() or not vectors_path.exists():
            raise RAGEngineError(f"Η συλλογή '{collection_name}' δεν είναι έτοιμη για αναζήτηση.")

        chunk_items = json.loads(chunks_path.read_text(encoding="utf-8"))
        vectors = np.load(vectors_path)

        query_vector = np.array(client.create_embeddings([query], model=embedding_model)[0], dtype=np.float32)
        query_vector = self.normalize(query_vector.reshape(1, -1))[0]
        scores = vectors @ query_vector
        if scores.ndim != 1:
            scores = scores.reshape(-1)

        top_indices = np.argsort(scores)[::-1][: max(1, top_k)]
        results: list[SearchResult] = []
        for idx in top_indices:
            item = chunk_items[int(idx)]
            results.append(
                SearchResult(
                    score=float(scores[int(idx)]),
                    filename=item["filename"],
                    chunk_id=item["chunk_id"],
                    chunk_index=int(item["chunk_index"]),
                    text=item["text"],
                )
            )
        return results

    @staticmethod
    def normalize(vectors: np.ndarray) -> np.ndarray:
        norms = np.linalg.norm(vectors, axis=1, keepdims=True)
        norms[norms == 0.0] = 1.0
        return vectors / norms


# ============================================================
# Prompt builder / auto-detect / fallback
# ============================================================
def resolve_model_identifier(
    models: list[AvailableModel],
    model_type: str,
    requested_identifier: str | None,
    preferred_identifier: str | None = None,
) -> str:
    filtered = [model for model in models if model.type == model_type]
    if not filtered:
        if model_type == "embedding":
            raise OllamaError("Δεν βρέθηκε διαθέσιμο embedding model στο Ollama. Κάνε pull ένα embedding model, π.χ. embeddinggemma ή qwen3-embedding.")
        raise OllamaError("Δεν βρέθηκε διαθέσιμο LLM model στο Ollama. Κάνε pull ένα chat/completion model από το Ollama.")

    requested_identifier = (requested_identifier or "").strip()
    preferred_identifier = (preferred_identifier or "").strip()

    if requested_identifier and requested_identifier != AUTO_OPTION:
        for model in filtered:
            if model.identifier == requested_identifier:
                return model.identifier
        raise OllamaError(f"Το επιλεγμένο model '{requested_identifier}' δεν βρέθηκε.")

    if preferred_identifier and preferred_identifier != AUTO_OPTION:
        for model in filtered:
            if model.identifier == preferred_identifier:
                return model.identifier

    loaded = [model for model in filtered if model.is_loaded]
    if loaded:
        loaded.sort(key=lambda m: (m.display_name or m.identifier).lower())
        return loaded[0].identifier

    if model_type == "embedding":
        hinted = [
            model
            for model in filtered
            if any(hint in (model.identifier + " " + (model.display_name or "")).lower() for hint in EMBEDDING_NAME_HINTS)
        ]
        if hinted:
            hinted.sort(key=lambda m: (m.display_name or m.identifier).lower())
            return hinted[0].identifier

    filtered.sort(key=lambda m: (m.display_name or m.identifier).lower())
    return filtered[0].identifier


def build_attachment_context(documents: list[LoadedDocument], max_chars: int = MAX_ATTACHMENT_CHARS) -> str:
    if not documents:
        return ""

    rendered_parts: list[str] = []
    total_chars = 0
    for doc in documents:
        block = (
            f"[ΕΠΙΣΥΝΑΠΤΟΜΕΝΟ ΑΡΧΕΙΟ]\n"
            f"Όνομα: {doc.filename}\n"
            f"Τύπος: {Path(doc.filename).suffix.lower() or 'unknown'}\n\n"
            f"Περιεχόμενο:\n{doc.text}\n"
        )
        if total_chars + len(block) > max_chars:
            remaining = max_chars - total_chars
            if remaining > 300:
                rendered_parts.append(truncate_middle(block, remaining))
            break
        rendered_parts.append(block)
        total_chars += len(block)

    return "\n\n".join(rendered_parts)


def build_chat_messages(
    question: str,
    rag_hits: list[SearchResult],
    attached_docs: list[LoadedDocument],
    task_mode: str,
) -> list[dict[str, str]]:
    source_blocks: list[str] = []
    for idx, hit in enumerate(rag_hits, start=1):
        source_blocks.append(
            f"[ΠΗΓΗ {idx}] αρχείο={hit.filename} | chunk={hit.chunk_index} | score={hit.score:.4f}\n{hit.text}"
        )

    rag_context = "\n\n".join(source_blocks)
    attachment_context = build_attachment_context(attached_docs)

    system_prompt = (
        "Είσαι τοπικός βοηθός RAG που δουλεύει πάνω σε context από βάση γνώσης και επισυναπτόμενα αρχεία. "
        "Να απαντάς πάντα στα Ελληνικά με φυσικό, σωστό και καθαρό λόγο, εκτός αν ο χρήστης ζητήσει ρητά άλλη γλώσσα. "
        "Να μη φαντάζεσαι facts που δεν υπάρχουν στο διαθέσιμο υλικό. "
        "Αν έχεις RAG πηγές, να χρησιμοποιείς παραπομπές [ΠΗΓΗ x]. "
        "Αν ο χρήστης ζητήσει έλεγχο ή διόρθωση κώδικα από συνημμένο αρχείο, να εντοπίζεις τα προβλήματα, "
        "να εξηγείς συνοπτικά τα σφάλματα και όταν χρειάζεται να επιστρέφεις ολόκληρο τον διορθωμένο κώδικα μέσα σε fenced code block. "
        "Αν το παρεχόμενο context δεν αρκεί, να το λες καθαρά."
    )

    mode_instruction = ""
    if task_mode == "Έλεγχος κώδικα":
        mode_instruction = (
            "Εστίασε σε code review. Βρες bugs, ασυνέπειες, πιθανά runtime προβλήματα, κακή δομή και βελτιώσεις."
        )
    elif task_mode == "Διόρθωση κώδικα":
        mode_instruction = (
            "Εστίασε στη διόρθωση. Εξήγησε σύντομα τι άλλαξες και δώσε ολόκληρο τον διορθωμένο κώδικα "
            "σε code block όπου είναι εφικτό."
        )
    elif task_mode == "Εξήγηση αρχείων":
        mode_instruction = "Εξήγησε το περιεχόμενο των επισυναπτόμενων αρχείων με απλό και δομημένο τρόπο."
    else:
        mode_instruction = "Απάντησε αξιοποιώντας πρώτα το διαθέσιμο context και μετά γενική λογική όπου επιτρέπεται."

    user_parts = [
        f"Task mode: {task_mode}",
        f"Οδηγία mode: {mode_instruction}",
    ]

    if rag_context:
        user_parts.append(f"RAG Context:\n\n{rag_context}")
    if attachment_context:
        user_parts.append(f"Επισυναπτόμενα αρχεία:\n\n{attachment_context}")

    user_parts.append(
        "Οδηγίες απάντησης:\n"
        "1. Να απαντάς στα Ελληνικά, με σωστούς τόνους και πλήρεις προτάσεις.\n"
        "2. Αν χρησιμοποιείς RAG πηγές, βάλε [ΠΗΓΗ x] μέσα στην απάντηση.\n"
        "3. Αν διορθώνεις κώδικα, δώσε πλήρη corrected version όταν είναι λογικό.\n"
        "4. Αν λείπουν πληροφορίες, δήλωσέ το καθαρά.\n"
        "5. Όταν υπάρχουν συνημμένα αρχεία, να τα θεωρείς πρωτεύον context για την ερώτηση."
    )
    user_parts.append(f"Ερώτηση χρήστη:\n{question}")

    return [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": "\n\n".join(user_parts)},
    ]


# ============================================================
# Worker infrastructure
# ============================================================
class WorkerSignals(QObject):
    result = Signal(object)
    error = Signal(object)
    progress = Signal(str)
    finished = Signal()

    @staticmethod
    def _safe_emit(signal_obj: Any, *args: Any) -> bool:
        try:
            signal_obj.emit(*args)
            return True
        except RuntimeError:
            return False

    def emit_result(self, payload: Any) -> bool:
        return self._safe_emit(self.result, payload)

    def emit_error(self, payload: Any) -> bool:
        return self._safe_emit(self.error, payload)

    def emit_progress(self, message: str) -> bool:
        return self._safe_emit(self.progress, message)

    def emit_finished(self) -> bool:
        return self._safe_emit(self.finished)


class FunctionWorker(QRunnable):
    def __init__(self, fn: Callable[..., Any], *args: Any, **kwargs: Any) -> None:
        super().__init__()
        self.fn = fn
        self.args = args
        self.kwargs = kwargs
        self.signals = WorkerSignals()

    @Slot()
    def run(self) -> None:
        try:
            if "progress_callback" not in self.kwargs:
                self.kwargs["progress_callback"] = self.signals.emit_progress
            result = self.fn(*self.args, **self.kwargs)
            self.signals.emit_result(result)
        except Exception as exc:
            self.signals.emit_error({"message": str(exc), "traceback": traceback.format_exc()})
        finally:
            self.signals.emit_finished()


# ============================================================
# Κύριο παράθυρο εφαρμογής
# ============================================================
class MainWindow(QMainWindow):
    def __init__(self) -> None:
        super().__init__()
        self.setWindowTitle(f"{APP_NAME} v{APP_VERSION}")
        self.resize(1380, 820)
        self.setMinimumSize(860, 560)

        self.settings = load_settings()
        self.session_log_file = LOGS_DIR / f"session_{datetime.now().strftime('%Y%m%d_%H%M%S')}.log"
        self.controller = OllamaController(self.settings)
        self.client = self.controller.client()
        self.engine = RAGEngine()
        self.thread_pool = QThreadPool.globalInstance()
        self.models_cache: list[AvailableModel] = []
        self.kb_file_paths: list[str] = []
        self.chat_attachment_paths: list[str] = []
        self.latest_answer: str = ""
        self.status_check_in_progress = False
        self._log_buffer: list[str] = []
        self._last_progress_message = ""
        self._last_progress_time = 0.0
        self.current_theme = "dark"
        self._synchronizing_model_combos = False
        self.models_operation_in_progress = False
        self._progress_lock_count = 0
        self._full_ui_busy = False
        self._is_shutting_down = False

        self._build_ui()
        self._apply_settings_to_ui()
        self._relax_window_geometry_constraints()
        self.refresh_collections()
        self.setup_startup_behaviour()

    # --------------------------------------------------------
    # UI build
    # --------------------------------------------------------
    def _build_ui(self) -> None:
        self._build_actions()

        root = QWidget()
        root.setMinimumSize(0, 0)
        self.setCentralWidget(root)
        main_layout = QVBoxLayout(root)
        main_layout.setContentsMargins(10, 10, 10, 10)
        main_layout.setSpacing(8)

        header = self._build_header()
        main_layout.addWidget(header)

        self.tabs = QTabWidget()
        self.tabs.setDocumentMode(True)
        self.tabs.addTab(self._build_connection_tab(), "🦙 Ollama")
        self.tabs.addTab(self._build_knowledge_tab(), "📚 Βάση Γνώσης")
        self.tabs.addTab(self._build_chat_tab(), "💬 Συνομιλία")
        self.tabs.currentChanged.connect(self.on_context_for_model_loading_changed)
        main_layout.addWidget(self.tabs, stretch=1)

        log_group = QGroupBox("📝 Καταγραφή ενεργειών")
        log_layout = QVBoxLayout(log_group)
        self.log_output = QPlainTextEdit()
        self.log_output.setReadOnly(True)
        self.log_output.setMinimumHeight(90)
        self.log_output.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        self.log_output.setMaximumBlockCount(1500)
        log_layout.addWidget(self.log_output)
        main_layout.addWidget(log_group)

        self.progress_bar = QProgressBar()
        self.progress_bar.setTextVisible(False)
        self.progress_bar.setRange(0, 1)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(False)
        main_layout.addWidget(self.progress_bar)

        self.status = QStatusBar()
        self.setStatusBar(self.status)
        self.status.showMessage("Έτοιμο")

        self.log_flush_timer = QTimer(self)
        self.log_flush_timer.timeout.connect(self.flush_logs)
        self.log_flush_timer.start(120)

        self.status_timer = QTimer(self)
        self.status_timer.timeout.connect(self.async_refresh_runtime_status)
        self.status_timer.start(8000)

        self.progress_watchdog_timer = QTimer(self)
        self.progress_watchdog_timer.timeout.connect(self._watch_progress_state)
        self.progress_watchdog_timer.start(250)

        self.apply_theme(self.current_theme)
        self._sanitize_widget_fonts()

    def _build_actions(self) -> None:
        menu = self.menuBar().addMenu("Αρχείο")

        act_save_settings = QAction("Αποθήκευση ρυθμίσεων", self)
        act_save_settings.triggered.connect(self.save_current_settings)
        menu.addAction(act_save_settings)

        act_exit = QAction("Έξοδος", self)
        act_exit.triggered.connect(self.close)
        menu.addAction(act_exit)

    def _build_header(self) -> QWidget:
        frame = QFrame()
        frame.setObjectName("HeaderFrame")
        layout = QHBoxLayout(frame)
        layout.setContentsMargins(16, 16, 16, 16)
        layout.setSpacing(12)

        text_col = QVBoxLayout()
        text_col.setSpacing(6)

        title = QLabel(APP_NAME)
        title.setObjectName("HeaderTitle")

        subtitle = QLabel(
            "Desktop εφαρμογή για Ollama με RAG, απλή συνομιλία χωρίς RAG, επισυναπτόμενα αρχεία και εργαλεία code review."
        )
        subtitle.setObjectName("HeaderSubtitle")
        subtitle.setWordWrap(True)

        badge = QLabel("Τοπικό RAG • Chat mode • Ollama")
        badge.setObjectName("HeaderBadge")

        hint = QLabel("Επίλεξε τρόπο συνομιλίας και το περιβάλλον θα φιλτράρει αυτόματα τα σχετικά μοντέλα και τα κατάλληλα πεδία.")
        hint.setObjectName("HeaderHint")
        hint.setWordWrap(True)

        text_col.addWidget(title)
        text_col.addWidget(subtitle)
        text_col.addWidget(badge, alignment=Qt.AlignmentFlag.AlignLeft)
        text_col.addWidget(hint)

        layout.addLayout(text_col, stretch=1)

        self.theme_btn = QToolButton()
        self.theme_btn.setObjectName("ThemeButton")
        self.theme_btn.clicked.connect(self.toggle_theme)
        layout.addWidget(self.theme_btn, alignment=Qt.AlignmentFlag.AlignTop)
        return frame

    def _build_connection_tab(self) -> QWidget:
        page = QWidget()
        page.setMinimumSize(0, 0)
        layout = QVBoxLayout(page)
        layout.setSpacing(10)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setChildrenCollapsible(False)
        splitter.addWidget(self._build_connection_settings_group())
        splitter.addWidget(self._build_connection_models_group())
        splitter.setSizes([620, 620])
        layout.addWidget(splitter)
        return self._wrap_in_scroll_area(page)

    def _build_connection_settings_group(self) -> QWidget:
        group = QGroupBox("Ρυθμίσεις σύνδεσης και εκκίνησης")
        layout = QVBoxLayout(group)

        form = QFormLayout()
        self.base_url_edit = QLineEdit()
        self.api_key_edit = QLineEdit()
        self.timeout_spin = QSpinBox()
        self.timeout_spin.setRange(5, 1200)

        self.exe_path_edit = QLineEdit()
        self.cli_path_edit = QLineEdit()

        exe_row = QHBoxLayout()
        exe_row.addWidget(self.exe_path_edit, stretch=1)
        self.browse_exe_btn = QPushButton("Browse...")
        self.browse_exe_btn.clicked.connect(self.select_exe_path)
        exe_row.addWidget(self.browse_exe_btn)

        cli_row = QHBoxLayout()
        cli_row.addWidget(self.cli_path_edit, stretch=1)
        self.browse_cli_btn = QPushButton("Browse...")
        self.browse_cli_btn.clicked.connect(self.select_cli_path)
        cli_row.addWidget(self.browse_cli_btn)

        form.addRow("Base URL", self.base_url_edit)
        form.addRow("API key", self.api_key_edit)
        form.addRow("Timeout (sec)", self.timeout_spin)
        form.addRow("Ollama executable", self._wrap_layout(exe_row))
        form.addRow("ollama CLI path", self._wrap_layout(cli_row))
        layout.addLayout(form)

        options_box = QGroupBox("Αυτόματη συμπεριφορά")
        options_layout = QVBoxLayout(options_box)
        self.auto_open_check = QCheckBox("Άνοιγμα Ollama αν δεν είναι ήδη ανοιχτό")
        self.auto_start_service_check = QCheckBox("Αυτόματη εκκίνηση service αν δεν απαντά")
        self.auto_refresh_models_check = QCheckBox("Αυτόματη ανανέωση μοντέλων στην εκκίνηση")
        self.auto_load_models_check = QCheckBox("Αυτόματο load των επιλεγμένων μοντέλων όταν χρειάζεται")
        options_layout.addWidget(self.auto_open_check)
        options_layout.addWidget(self.auto_start_service_check)
        options_layout.addWidget(self.auto_refresh_models_check)
        options_layout.addWidget(self.auto_load_models_check)
        layout.addWidget(options_box)

        status_box = QGroupBox("Κατάσταση Ollama")
        status_layout = QFormLayout(status_box)
        self.app_status_label = QLabel("Άγνωστο")
        self.server_status_label = QLabel("Άγνωστο")
        self.models_status_label = QLabel("Δεν έχουν φορτωθεί μοντέλα")
        status_layout.addRow("App", self.app_status_label)
        status_layout.addRow("Service", self.server_status_label)
        status_layout.addRow("Models", self.models_status_label)
        layout.addWidget(status_box)

        button_row = QGridLayout()
        self.save_settings_btn = QPushButton("💾 Αποθήκευση ρυθμίσεων")
        self.save_settings_btn.clicked.connect(self.save_current_settings)
        self.detect_paths_btn = QPushButton("🔎 Εύρεση paths")
        self.detect_paths_btn.clicked.connect(self.detect_paths)
        self.check_status_btn = QPushButton("📡 Έλεγχος κατάστασης")
        self.check_status_btn.clicked.connect(self.refresh_runtime_status)
        self.open_app_btn = QPushButton("🚀 Άνοιγμα Ollama")
        self.open_app_btn.clicked.connect(self.open_ollama_clicked)
        self.start_service_btn = QPushButton("🌐 Εκκίνηση Ollama Service")
        self.start_service_btn.clicked.connect(self.start_service_clicked)
        self.refresh_models_btn = QPushButton("🔄 Ανανέωση μοντέλων")
        self.refresh_models_btn.clicked.connect(self.refresh_models_clicked)

        button_row.addWidget(self.save_settings_btn, 0, 0)
        button_row.addWidget(self.detect_paths_btn, 0, 1)
        button_row.addWidget(self.check_status_btn, 0, 2)
        button_row.addWidget(self.open_app_btn, 1, 0)
        button_row.addWidget(self.start_service_btn, 1, 1)
        button_row.addWidget(self.refresh_models_btn, 1, 2)
        layout.addLayout(button_row)
        layout.addStretch(1)
        return group

    def _build_connection_models_group(self) -> QWidget:
        group = QGroupBox("Μοντέλα και auto-detect")
        layout = QVBoxLayout(group)

        info = QLabel(
            "Η εφαρμογή διαχωρίζει αυτόματα τα LLM από τα embedding models. "
            "Στη συνομιλία χωρίς RAG δίνεται έμφαση μόνο στα chat μοντέλα, ενώ στο RAG προβάλλονται και τα embeddings."
        )
        info.setWordWrap(True)
        layout.addWidget(info)

        form = QFormLayout()
        self.llm_combo = QComboBox()
        self.embedding_combo = QComboBox()
        self.models_filter_combo = QComboBox()
        self.model_load_scope_combo = QComboBox()
        self.llm_combo.setMinimumWidth(260)
        self.embedding_combo.setMinimumWidth(260)
        self.models_filter_combo.addItem("Με βάση τη λειτουργία συνομιλίας", MODEL_LIST_FILTER_MODE)
        self.models_filter_combo.addItem("Μόνο LLM", MODEL_LIST_FILTER_LLM)
        self.models_filter_combo.addItem("Μόνο Embedding", MODEL_LIST_FILTER_EMBEDDING)
        self.models_filter_combo.addItem("Όλα τα μοντέλα", MODEL_LIST_FILTER_ALL)
        self.models_filter_combo.currentIndexChanged.connect(self.update_models_list_display)
        self.model_load_scope_combo.addItem("Μόνο το ενεργό μοντέλο για το τρέχον context", MODEL_LOAD_SCOPE_ACTIVE)
        self.model_load_scope_combo.addItem("Μόνο LLM", MODEL_LOAD_SCOPE_LLM)
        self.model_load_scope_combo.addItem("Μόνο Embedding", MODEL_LOAD_SCOPE_EMBEDDING)
        self.model_load_scope_combo.addItem("Και τα δύο", MODEL_LOAD_SCOPE_BOTH)
        self.model_load_scope_combo.currentIndexChanged.connect(self.on_context_for_model_loading_changed)
        self.llm_combo.currentIndexChanged.connect(self.on_model_selection_changed)
        self.embedding_combo.currentIndexChanged.connect(self.on_model_selection_changed)
        form.addRow("Προεπιλεγμένο LLM", self.llm_combo)
        form.addRow("Προεπιλεγμένο Embedding", self.embedding_combo)
        form.addRow("Προβολή λίστας", self.models_filter_combo)
        form.addRow("Τι να φορτώνεται", self.model_load_scope_combo)
        layout.addLayout(form)

        self.selected_models_summary = QLabel("Δεν έχουν ανανεωθεί ακόμη μοντέλα.")
        self.selected_models_summary.setObjectName("ModeHintLabel")
        self.selected_models_summary.setWordWrap(True)
        layout.addWidget(self.selected_models_summary)

        self.model_search_edit = QLineEdit()
        self.model_search_edit.setPlaceholderText("Φιλτράρισμα μοντέλων με όνομα ή publisher...")
        self.model_search_edit.textChanged.connect(self.update_models_list_display)
        layout.addWidget(self.model_search_edit)

        self.models_list = QListWidget()
        self.models_list.setObjectName("ModelsList")
        self.models_list.setMinimumHeight(96)
        self.models_list.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        layout.addWidget(self.models_list)

        button_row = QHBoxLayout()
        self.load_models_btn = QPushButton("📥 Φόρτωση ενεργού μοντέλου")
        self.load_models_btn.clicked.connect(self.load_selected_models_clicked)
        self.resolve_models_btn = QPushButton("🤖 Δοκιμή auto-detect")
        self.resolve_models_btn.clicked.connect(self.resolve_models_preview)
        button_row.addWidget(self.load_models_btn)
        button_row.addWidget(self.resolve_models_btn)
        layout.addLayout(button_row)
        layout.addStretch(1)
        return group

    def _build_knowledge_tab(self) -> QWidget:
        page = QWidget()
        page.setMinimumSize(0, 0)
        layout = QVBoxLayout(page)

        splitter = QSplitter(Qt.Orientation.Horizontal)
        splitter.setChildrenCollapsible(False)
        splitter.addWidget(self._build_kb_build_group())
        splitter.addWidget(self._build_kb_collections_group())
        splitter.setSizes([700, 520])
        layout.addWidget(splitter)
        return self._wrap_in_scroll_area(page)

    def _build_kb_build_group(self) -> QWidget:
        group = QGroupBox("Δημιουργία συλλογής γνώσης")
        layout = QVBoxLayout(group)

        form = QFormLayout()
        self.collection_name_edit = QLineEdit()
        self.chunk_size_spin = QSpinBox()
        self.chunk_size_spin.setRange(100, 12000)
        self.chunk_overlap_spin = QSpinBox()
        self.chunk_overlap_spin.setRange(0, 6000)
        self.embed_batch_spin = QSpinBox()
        self.embed_batch_spin.setRange(1, 256)
        form.addRow("Όνομα συλλογής", self.collection_name_edit)
        form.addRow("Chunk size", self.chunk_size_spin)
        form.addRow("Chunk overlap", self.chunk_overlap_spin)
        form.addRow("Batch size embeddings", self.embed_batch_spin)
        layout.addLayout(form)

        row = QHBoxLayout()
        self.select_kb_files_btn = QPushButton("📂 Επιλογή αρχείων")
        self.select_kb_files_btn.clicked.connect(self.select_kb_files)
        self.clear_kb_files_btn = QPushButton("🧹 Καθαρισμός λίστας")
        self.clear_kb_files_btn.clicked.connect(self.clear_kb_files)
        row.addWidget(self.select_kb_files_btn)
        row.addWidget(self.clear_kb_files_btn)
        layout.addLayout(row)

        self.kb_files_list = QListWidget()
        self.kb_files_list.setMinimumHeight(88)
        self.kb_files_list.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        layout.addWidget(self.kb_files_list)

        self.build_collection_btn = QPushButton("🏗️ Δημιουργία / ενημέρωση συλλογής")
        self.build_collection_btn.clicked.connect(self.build_collection_clicked)
        layout.addWidget(self.build_collection_btn)
        layout.addStretch(1)
        return group

    def _build_kb_collections_group(self) -> QWidget:
        group = QGroupBox("Υπάρχουσες συλλογές")
        layout = QVBoxLayout(group)

        row = QHBoxLayout()
        self.refresh_collections_btn = QPushButton("🔄 Ανανέωση")
        self.refresh_collections_btn.clicked.connect(self.refresh_collections)
        self.delete_collection_btn = QPushButton("🗑️ Διαγραφή")
        self.delete_collection_btn.clicked.connect(self.delete_selected_collection)
        row.addWidget(self.refresh_collections_btn)
        row.addWidget(self.delete_collection_btn)
        layout.addLayout(row)

        self.collections_list = QListWidget()
        self.collections_list.currentItemChanged.connect(self.show_selected_collection_metadata)
        layout.addWidget(self.collections_list)

        self.collection_metadata_output = QPlainTextEdit()
        self.collection_metadata_output.setReadOnly(True)
        self.collection_metadata_output.setPlaceholderText("Τα μεταδεδομένα της επιλεγμένης συλλογής θα εμφανιστούν εδώ.")
        self.collection_metadata_output.setMinimumHeight(82)
        self.collection_metadata_output.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        layout.addWidget(self.collection_metadata_output)
        return group

    def _build_chat_tab(self) -> QWidget:
        page = QWidget()
        page.setMinimumSize(0, 0)
        layout = QVBoxLayout(page)

        upper = QSplitter(Qt.Orientation.Horizontal)
        upper.setChildrenCollapsible(False)
        upper.addWidget(self._build_chat_controls_group())
        upper.addWidget(self._build_chat_output_group())
        upper.setSizes([560, 900])
        layout.addWidget(upper)
        return self._wrap_in_scroll_area(page)

    def _build_chat_controls_group(self) -> QWidget:
        group = QGroupBox("Ρυθμίσεις συνομιλίας")
        layout = QVBoxLayout(group)

        form = QFormLayout()
        self.chat_mode_combo = QComboBox()
        self.chat_mode_combo.addItem(CHAT_MODE_PLAIN, "plain")
        self.chat_mode_combo.addItem(CHAT_MODE_RAG, "rag")
        self.chat_mode_combo.currentIndexChanged.connect(self.on_chat_mode_changed)
        self.chat_collection_combo = QComboBox()
        self.task_mode_combo = QComboBox()
        self.task_mode_combo.addItems([
            "Αυτόματο",
            "Γενική ερώτηση",
            "Έλεγχος κώδικα",
            "Διόρθωση κώδικα",
            "Εξήγηση αρχείων",
        ])
        self.top_k_spin = QSpinBox()
        self.top_k_spin.setRange(1, 20)
        self.temperature_spin = QDoubleSpinBox()
        self.temperature_spin.setRange(0.0, 2.0)
        self.temperature_spin.setSingleStep(0.1)
        self.max_tokens_spin = QSpinBox()
        self.max_tokens_spin.setRange(64, 16384)
        form.addRow("Λειτουργία", self.chat_mode_combo)
        form.addRow("Συλλογή γνώσης", self.chat_collection_combo)
        form.addRow("Task mode", self.task_mode_combo)
        form.addRow("Top-K retrieval", self.top_k_spin)
        form.addRow("Temperature", self.temperature_spin)
        form.addRow("Max tokens", self.max_tokens_spin)
        layout.addLayout(form)

        self.chat_mode_hint_label = QLabel()
        self.chat_mode_hint_label.setObjectName("ModeHintLabel")
        self.chat_mode_hint_label.setWordWrap(True)
        layout.addWidget(self.chat_mode_hint_label)

        attach_box = QGroupBox("Επισυναπτόμενα αρχεία στην ερώτηση")
        attach_layout = QVBoxLayout(attach_box)
        row = QHBoxLayout()
        self.add_chat_files_btn = QPushButton("📎 Προσθήκη αρχείων")
        self.add_chat_files_btn.clicked.connect(self.select_chat_attachments)
        self.clear_chat_files_btn = QPushButton("🧹 Καθαρισμός")
        self.clear_chat_files_btn.clicked.connect(self.clear_chat_attachments)
        row.addWidget(self.add_chat_files_btn)
        row.addWidget(self.clear_chat_files_btn)
        attach_layout.addLayout(row)
        self.chat_files_list = QListWidget()
        self.chat_files_list.setMinimumHeight(82)
        self.chat_files_list.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        attach_layout.addWidget(self.chat_files_list)
        layout.addWidget(attach_box)

        question_label = QLabel("Ερώτηση")
        layout.addWidget(question_label)
        self.question_input = QPlainTextEdit()
        self.question_input.setPlaceholderText(
            "Γράψε ερώτηση...\n\nΠαραδείγματα:\n- Εξήγησε αυτά τα αρχεία\n- Έλεγξε τον κώδικα και βρες bugs\n- Διόρθωσε το app.py και δώσε μου ολόκληρο τον corrected code"
        )
        self.question_input.setMinimumHeight(128)
        self.question_input.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        layout.addWidget(self.question_input)

        button_row = QHBoxLayout()
        self.ask_btn = QPushButton("💬 Υποβολή ερώτησης")
        self.ask_btn.clicked.connect(self.ask_question_clicked)
        self.clear_answer_btn = QPushButton("🧼 Καθαρισμός απάντησης")
        self.clear_answer_btn.clicked.connect(self.clear_answer_output)
        button_row.addWidget(self.ask_btn)
        button_row.addWidget(self.clear_answer_btn)
        layout.addLayout(button_row)
        layout.addStretch(1)
        return group

    def _build_chat_output_group(self) -> QWidget:
        group = QGroupBox("Απάντηση και πηγές")
        layout = QVBoxLayout(group)

        self.answer_output = QTextEdit()
        self.answer_output.setObjectName("AnswerOutput")
        self.answer_output.setReadOnly(True)
        self.answer_output.setAcceptRichText(True)
        self.answer_output.setLineWrapMode(QTextEdit.LineWrapMode.WidgetWidth)
        self.answer_output.setWordWrapMode(QTextOption.WrapAtWordBoundaryOrAnywhere)
        self.answer_output.setHorizontalScrollBarPolicy(Qt.ScrollBarPolicy.ScrollBarAlwaysOff)
        self.answer_output.document().setDocumentMargin(12)
        self.answer_output.setPlaceholderText("Η απάντηση θα εμφανιστεί εδώ σε μορφοποιημένο, κυλιόμενο πλαίσιο.")
        self.answer_output.setMinimumHeight(220)
        self.answer_output.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        layout.addWidget(self.answer_output, stretch=1)

        save_row = QHBoxLayout()
        self.copy_answer_btn = QPushButton("📋 Αντιγραφή απάντησης")
        self.copy_answer_btn.clicked.connect(self.copy_answer_to_clipboard)
        self.save_answer_btn = QPushButton("💾 Αποθήκευση απάντησης")
        self.save_answer_btn.clicked.connect(self.save_answer_text)
        self.save_code_btn = QPushButton("🧩 Αποθήκευση corrected code")
        self.save_code_btn.clicked.connect(self.save_corrected_code)
        save_row.addWidget(self.copy_answer_btn)
        save_row.addWidget(self.save_answer_btn)
        save_row.addWidget(self.save_code_btn)
        layout.addLayout(save_row)

        self.sources_title_label = QLabel("Πηγές RAG / συνημμένα")
        layout.addWidget(self.sources_title_label)
        self.sources_output = QPlainTextEdit()
        self.sources_output.setObjectName("SourcesOutput")
        self.sources_output.setReadOnly(True)
        self.sources_output.setPlaceholderText("Εδώ θα εμφανιστούν οι πηγές RAG ή τα επισυναπτόμενα αρχεία που χρησιμοποιήθηκαν.")
        self.sources_output.setMinimumHeight(76)
        self.sources_output.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
        layout.addWidget(self.sources_output)
        return group

    # --------------------------------------------------------
    # Helper UI
    # --------------------------------------------------------
    def _wrap_layout(self, layout: QHBoxLayout) -> QWidget:
        widget = QWidget()
        widget.setLayout(layout)
        return widget

    def _wrap_in_scroll_area(self, content: QWidget) -> QScrollArea:
        area = QScrollArea()
        area.setWidgetResizable(True)
        area.setFrameShape(QFrame.Shape.NoFrame)
        area.setWidget(content)
        area.setMinimumSize(0, 0)
        content.setMinimumSize(0, 0)
        return area

    def _sanitize_widget_fonts(self) -> None:
        app = QApplication.instance()
        if app is None:
            return
        base_font = app.font()
        family = base_font.family() or "Segoe UI"
        point_size = base_font.pointSizeF()
        safe_size = max(11.0, point_size if point_size and point_size > 0 else 11.0)
        safe_font = QFont(family)
        safe_font.setPointSizeF(safe_size)
        for widget in [self, *self.findChildren(QWidget)]:
            try:
                current = widget.font()
                if current.pointSizeF() <= 0:
                    widget.setFont(safe_font)
            except Exception:
                pass

    def _relax_window_geometry_constraints(self) -> None:
        try:
            self.setMinimumSize(860, 560)
            central = self.centralWidget()
            if central is not None:
                central.setMinimumSize(0, 0)
            if hasattr(self, "tabs"):
                self.tabs.setMinimumSize(0, 0)
            for widget in (
                self.log_output,
                self.models_list,
                self.kb_files_list,
                self.collection_metadata_output,
                self.chat_files_list,
                self.question_input,
                self.answer_output,
                self.sources_output,
            ):
                widget.setMinimumHeight(max(56, min(widget.minimumHeight(), 160)))
                widget.setMinimumWidth(0)
                widget.setSizePolicy(QSizePolicy.Policy.Expanding, QSizePolicy.Policy.Expanding)
            self.answer_output.setMinimumHeight(220)
            self.question_input.setMinimumHeight(128)
            self.models_list.setMinimumHeight(96)
            self.sources_output.setMinimumHeight(92)
            self.log_output.setMinimumHeight(90)
        except Exception:
            pass

    def create_light_palette(self) -> QPalette:
        palette = QPalette()
        palette.setColor(QPalette.ColorRole.Window, QColor("#f0f4fc"))
        palette.setColor(QPalette.ColorRole.WindowText, QColor("#1a1d23"))
        palette.setColor(QPalette.ColorRole.Base, QColor("#ffffff"))
        palette.setColor(QPalette.ColorRole.AlternateBase, QColor("#f5f7fa"))
        palette.setColor(QPalette.ColorRole.ToolTipBase, QColor("#1a2a4a"))
        palette.setColor(QPalette.ColorRole.ToolTipText, QColor("#e8f0fe"))
        palette.setColor(QPalette.ColorRole.Text, QColor("#1a1d23"))
        palette.setColor(QPalette.ColorRole.PlaceholderText, QColor("#8a9ab0"))
        palette.setColor(QPalette.ColorRole.Button, QColor("#0ea5e9"))
        palette.setColor(QPalette.ColorRole.ButtonText, QColor("#ffffff"))
        palette.setColor(QPalette.ColorRole.Highlight, QColor("#0284c7"))
        palette.setColor(QPalette.ColorRole.HighlightedText, QColor("#ffffff"))
        palette.setColor(QPalette.ColorRole.Link, QColor("#1e88e5"))
        palette.setColor(QPalette.ColorRole.LinkVisited, QColor("#7b1fa2"))
        palette.setColor(QPalette.ColorRole.Mid, QColor("#c8d4e8"))
        palette.setColor(QPalette.ColorRole.Midlight, QColor("#dde6f5"))
        palette.setColor(QPalette.ColorRole.Dark, QColor("#4a5568"))
        palette.setColor(QPalette.ColorRole.Shadow, QColor("#1a1d23"))
        return palette

    def _dark_stylesheet(self) -> str:
        return """
        QWidget {
            background-color: #0b1018;
            color: #dce6f8;
            font-family: "Segoe UI", "Inter", Arial, sans-serif;
            font-size: 14px;
        }
        QMainWindow { background-color: #0a0e18; }
        QDialog { background-color: #0d1220; }
        QLabel { background: transparent; color: #dce6f8; border: none; }

        QMenuBar {
            background: #0a0e18;
            color: #dce6f8;
            border-bottom: 1px solid #1f2a44;
        }
        QMenuBar::item:selected { background: #162038; color: #00d4ff; border-radius: 6px; }
        QMenu {
            background-color: #0d1628;
            color: #dce6f8;
            border: 1.5px solid #1e2d50;
            border-top: 2px solid #00d4ff;
            border-radius: 10px;
            padding: 5px;
        }
        QMenu::item { padding: 7px 30px 7px 16px; border-radius: 6px; }
        QMenu::item:selected {
            background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #00d4ff, stop:1 #0284c7);
            color: #06101c;
        }

        #HeaderFrame {
            background: qlineargradient(x1:0, y1:0, x2:1, y2:0,
                                        stop:0 #0d1117,
                                        stop:0.35 #0f1922,
                                        stop:0.7 #122238,
                                        stop:1 #0d1117);
            border: 1.5px solid #1f2a44;
            border-radius: 18px;
        }
        #HeaderTitle { color: #f0f6fc; }
        #HeaderSubtitle { color: #9fb7d9; }
        #HeaderBadge {
            color: #b3e6ff;
            background: rgba(0, 212, 255, 0.12);
            border: 1px solid rgba(0, 212, 255, 0.35);
            border-radius: 10px;
            padding: 6px 10px;
            font-weight: 600;
        }
        #HeaderHint { color: #8fb4df; font-size: 12px; padding-right: 4px; }
        QToolButton#ThemeButton {
            background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 #111827, stop:1 #172640);
            color: #e6f7ff;
            border: 1.5px solid #00d4ff;
            border-radius: 12px;
            padding: 9px 16px;
            font-weight: 700;
        }
        QToolButton#ThemeButton:hover { background: #1d3152; }

        QGroupBox {
            background: #0d1628;
            border: 1px solid #1e2d50;
            border-top: 2px solid #00d4ff;
            border-radius: 16px;
            margin-top: 12px;
            font-weight: 700;
            padding-top: 16px;
        }
        QGroupBox::title {
            left: 14px;
            padding: 0 8px;
            color: #00d4ff;
        }
        QPushButton {
            background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #1f6feb, stop:1 #0d47a1);
            color: #ffffff;
            border: 1px solid #2d69d6;
            border-radius: 12px;
            padding: 10px 14px;
            font-weight: 700;
        }
        QPushButton:hover { background: #388bfd; }
        QPushButton:pressed { background: #1f6feb; }
        QPushButton:disabled {
            background: #1a2538;
            color: #6280a8;
            border-color: #23324b;
        }
        QLineEdit, QPlainTextEdit, QTextEdit, QComboBox, QListWidget, QSpinBox, QDoubleSpinBox {
            background-color: #111827;
            color: #dce6f8;
            border: 1px solid #1f2d45;
            border-radius: 10px;
            padding: 8px;
            selection-background-color: #1f6feb;
            selection-color: #ffffff;
        }
        QLineEdit:focus, QPlainTextEdit:focus, QTextEdit:focus, QComboBox:focus { border: 1.5px solid #00d4ff; }
        QListWidget::item { padding: 8px 6px; border-radius: 8px; }
        QListWidget::item:selected {
            background: #172640;
            color: #e6fbff;
            border: 1px solid #00d4ff;
        }
        QTabWidget::pane {
            border: none;
            border-top: 2.5px solid #00d4ff;
            border-radius: 14px;
            background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #0c1420, stop:1 #0d1220);
        }
        QTabBar { background: #0a0e18; qproperty-drawBase: 0; }
        QTabBar::tab {
            background: #11182a;
            color: #6888bb;
            padding: 9px 18px;
            min-height: 34px;
            border-top-left-radius: 9px;
            border-top-right-radius: 9px;
            margin-right: 3px;
            font-weight: 700;
            border: 1px solid #1e2d50;
            border-bottom: none;
        }
        QTabBar::tab:hover {
            background: #162038;
            color: #00d4ff;
            border-color: #00aacc;
        }
        QTabBar::tab:selected {
            background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #1e3050, stop:1 #172640);
            color: #00e5ff;
            border: 2px solid #00d4ff;
            border-bottom: 2px solid #172640;
            margin-bottom: -2px;
        }
        QStatusBar {
            background: #0d1628;
            color: #5a7aaa;
            border-top: 1.5px solid #00d4ff;
            font-size: 11.5pt;
        }
        QStatusBar::item { border: none; }
        QProgressBar {
            border: 1px solid #1e2d50;
            border-radius: 8px;
            background: #11182a;
        }
        QProgressBar::chunk {
            background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 #00d4ff, stop:1 #1f6feb);
            border-radius: 8px;
        }
        QSplitter::handle { background: #132037; }
        """

    def _light_stylesheet(self) -> str:
        return """
        QWidget {
            background-color: #f0f4fc;
            color: #1a1d23;
            font-family: "Segoe UI", "Inter", Arial, sans-serif;
            font-size: 14px;
        }
        QMainWindow { background-color: #eef3fb; }
        QDialog { background-color: #f5f8fd; }
        QLabel { background: transparent; color: #1a1d23; border: none; }

        QMenuBar {
            background: #eaf2fd;
            color: #12385a;
            border-bottom: 1px solid #c8d7ee;
        }
        QMenuBar::item:selected { background: #d5e8fb; color: #0284c7; border-radius: 6px; }
        QMenu {
            background-color: #ffffff;
            color: #0a0f1e;
            border: 1.5px solid #bdd6f5;
            border-top: 2px solid #0ea5e9;
            border-radius: 10px;
            padding: 5px;
        }
        QMenu::item { padding: 7px 30px 7px 16px; border-radius: 6px; }
        QMenu::item:selected {
            background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #0ea5e9, stop:1 #0284c7);
            color: #ffffff;
        }

        #HeaderFrame {
            background: qlineargradient(x1:0, y1:0, x2:1, y2:0, stop:0 #f8fbff, stop:0.45 #eef6ff, stop:1 #e7f1fd);
            border: 1.5px solid #c8d7ee;
            border-top: 3px solid #0ea5e9;
            border-radius: 18px;
        }
        #HeaderTitle { color: #12385a; font-size: 22px; font-weight: 800; }
        #HeaderSubtitle { color: #46607e; font-size: 13px; }
        #HeaderBadge {
            color: #045f8f;
            background: rgba(14, 165, 233, 0.10);
            border: 1px solid rgba(14, 165, 233, 0.28);
            border-radius: 10px;
            padding: 6px 10px;
            font-weight: 600;
        }
        #HeaderHint { color: #46607e; font-size: 12px; padding-right: 4px; }
        #ModeHintLabel { background: rgba(14, 165, 233, 0.09); border: 1px solid rgba(14, 165, 233, 0.25); border-radius: 12px; padding: 9px 11px; color: #24507b; }
        QTextEdit#AnswerOutput, QPlainTextEdit#SourcesOutput { border: 1px solid #a9bfdc; }
        QListWidget#ModelsList { border: 1px solid #a9bfdc; }
        QToolButton#ThemeButton {
            background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 #0ea5e9, stop:1 #0284c7);
            color: #ffffff;
            border: 1px solid #0284c7;
            border-radius: 12px;
            padding: 9px 16px;
            font-weight: 700;
        }

        QGroupBox {
            background: #ffffff;
            border: 1px solid #d8e2f0;
            border-top: 2px solid #0ea5e9;
            border-radius: 16px;
            margin-top: 12px;
            font-weight: 700;
            padding-top: 16px;
        }
        QGroupBox::title {
            left: 14px;
            padding: 0 8px;
            color: #0284c7;
        }
        QPushButton {
            background: qlineargradient(x1:0,y1:0,x2:0,y2:1, stop:0 #0ea5e9, stop:1 #0284c7);
            color: #ffffff;
            border: 1px solid #0284c7;
            border-radius: 12px;
            padding: 10px 14px;
            font-weight: 700;
        }
        QPushButton:hover { background: #38bdf8; }
        QPushButton:disabled {
            background: #b8d7ec;
            color: #ffffff;
            border-color: #a1c5e0;
        }
        QLineEdit, QPlainTextEdit, QTextEdit, QComboBox, QListWidget, QSpinBox, QDoubleSpinBox {
            background-color: #ffffff;
            color: #1a1d23;
            border: 1px solid #c8d4e8;
            border-radius: 10px;
            padding: 8px;
            selection-background-color: #0284c7;
            selection-color: #ffffff;
        }
        QLineEdit:focus, QPlainTextEdit:focus, QTextEdit:focus, QComboBox:focus { border: 1.5px solid #0ea5e9; }
        QListWidget::item { padding: 8px 6px; border-radius: 8px; }
        QListWidget::item:selected {
            background: #e0f2fe;
            color: #075985;
            border: 1px solid #38bdf8;
        }
        QTabWidget::pane {
            border: 1px solid #d9e2ef;
            border-top: 2.5px solid #0ea5e9;
            border-radius: 14px;
            background: #eef3fb;
        }
        QTabBar::tab {
            background: #dce6f7;
            color: #36557b;
            padding: 9px 18px;
            min-height: 34px;
            border-top-left-radius: 9px;
            border-top-right-radius: 9px;
            margin-right: 3px;
            font-weight: 700;
            border: 1px solid #c6d7ee;
            border-bottom: none;
        }
        QTabBar::tab:selected {
            background: #ffffff;
            color: #0284c7;
            border: 2px solid #0ea5e9;
            border-bottom: 2px solid #ffffff;
            margin-bottom: -2px;
        }
        QStatusBar {
            background: #ffffff;
            color: #1e3a5f;
            border-top: 1.5px solid #0ea5e9;
            font-size: 11.5pt;
        }
        QStatusBar::item { border: none; }
        QProgressBar {
            border: 1px solid #c8d4e8;
            border-radius: 8px;
            background: #ffffff;
        }
        QProgressBar::chunk {
            background: qlineargradient(x1:0,y1:0,x2:1,y2:0, stop:0 #0ea5e9, stop:1 #0284c7);
            border-radius: 8px;
        }
        QSplitter::handle { background: #d6e2f0; }
        """

    def apply_theme(self, mode: str) -> None:
        app = QApplication.instance()
        self.current_theme = "dark" if (mode or "").strip().lower() == "dark" else "light"
        if app is not None:
            if self.current_theme == "light":
                app.setPalette(self.create_light_palette())
            else:
                app.setPalette(app.style().standardPalette())
        self.setStyleSheet(self._dark_stylesheet() if self.current_theme == "dark" else self._light_stylesheet())
        if hasattr(self, "theme_btn"):
            self.theme_btn.setText("🌙 Σκούρο" if self.current_theme == "dark" else "☀️ Φωτεινό")

    def toggle_theme(self) -> None:
        self.apply_theme("light" if self.current_theme == "dark" else "dark")


    def log(self, message: str, immediate: bool = False) -> None:
        timestamp = datetime.now().strftime("%H:%M:%S")
        line = f"[{timestamp}] {message}"
        try:
            with self.session_log_file.open("a", encoding="utf-8") as fh:
                fh.write(line + "\n")
        except Exception:
            pass

        if self._is_shutting_down:
            return

        try:
            if immediate:
                self.log_output.appendPlainText(line)
                self.log_output.moveCursor(QTextCursor.MoveOperation.End)
            else:
                self._log_buffer.append(line)
            self.status.showMessage(message.splitlines()[0], 5000)
        except RuntimeError:
            return

    def log_exception_details(self, title: str, details: str) -> None:
        details = (details or "").strip()
        if not details:
            return
        block = f"{title}\n{details}"
        self.log(block, immediate=True)

    def flush_logs(self) -> None:
        if self._is_shutting_down or not self._log_buffer:
            return
        try:
            self.log_output.appendPlainText("\n".join(self._log_buffer))
            self._log_buffer.clear()
            self.log_output.moveCursor(QTextCursor.MoveOperation.End)
        except RuntimeError:
            self._log_buffer.clear()

    def on_worker_progress(self, message: str) -> None:
        if self._is_shutting_down:
            return

        now = time.monotonic()
        if message == self._last_progress_message and now - self._last_progress_time < 0.6:
            return

        if "Δημιουργία embeddings batch" in message:
            match = re.search(r"(\d+)/(\d+)", message)
            if match:
                current = int(match.group(1))
                total = int(match.group(2))
                if current not in {1, total} and total > 6 and current % max(1, total // 6) != 0:
                    if now - self._last_progress_time < 1.0:
                        return

        self._last_progress_message = message
        self._last_progress_time = now
        self.log(message)

    def show_error(self, message: str, details: str | None = None) -> None:
        self.log(f"❌ {message}")
        if details:
            self.log_exception_details("🧾 Αναλυτικά στοιχεία σφάλματος:", details)
        if self._is_shutting_down:
            return
        try:
            QMessageBox.critical(self, APP_NAME, message)
        except RuntimeError:
            return

    def show_info(self, message: str) -> None:
        self.log(f"ℹ️ {message}")
        if self._is_shutting_down:
            return
        try:
            QMessageBox.information(self, APP_NAME, message)
        except RuntimeError:
            return

    @staticmethod
    def _normalize_worker_error_payload(payload: Any) -> tuple[str, str]:
        if isinstance(payload, dict):
            message = str(payload.get("message") or "Άγνωστο σφάλμα")
            details = str(payload.get("traceback") or "").strip()
            return message, details
        return str(payload or "Άγνωστο σφάλμα"), ""

    def _set_main_controls_enabled(self, enabled: bool) -> None:
        for button in [
            self.save_settings_btn,
            self.detect_paths_btn,
            self.check_status_btn,
            self.open_app_btn,
            self.start_service_btn,
            self.refresh_models_btn,
            self.load_models_btn,
            self.resolve_models_btn,
            self.select_kb_files_btn,
            self.clear_kb_files_btn,
            self.build_collection_btn,
            self.refresh_collections_btn,
            self.delete_collection_btn,
            self.add_chat_files_btn,
            self.clear_chat_files_btn,
            self.ask_btn,
            self.clear_answer_btn,
            self.copy_answer_btn,
            self.save_answer_btn,
            self.save_code_btn,
        ]:
            button.setEnabled(enabled)

    def _set_models_controls_enabled(self, enabled: bool) -> None:
        for widget in [
            self.refresh_models_btn,
            self.load_models_btn,
            self.resolve_models_btn,
            self.llm_combo,
            self.embedding_combo,
            self.models_filter_combo,
            self.model_load_scope_combo,
            self.model_search_edit,
            self.models_list,
        ]:
            widget.setEnabled(enabled)

    def _hard_reset_progress_state(self, message: str | None = None) -> None:
        self._progress_lock_count = 0
        self._full_ui_busy = False
        self.models_operation_in_progress = False
        self.progress_bar.setRange(0, 1)
        self.progress_bar.setValue(0)
        self.progress_bar.setVisible(False)
        self.progress_bar.update()
        self._set_main_controls_enabled(True)
        self._set_models_controls_enabled(True)
        self.refresh_models_btn.setEnabled(True)
        self.load_models_btn.setEnabled(True)
        if message:
            self.status.showMessage(message, 5000)

    def _watch_progress_state(self) -> None:
        has_threads = self.thread_pool.activeThreadCount() > 0
        if has_threads:
            return
        if self.status_check_in_progress:
            return
        if self.progress_bar.isVisible() or self._progress_lock_count > 0 or self._full_ui_busy or self.models_operation_in_progress:
            self._hard_reset_progress_state("Έτοιμο")

    def start_progress_indicator(self, message: str | None = None) -> None:
        self._progress_lock_count += 1
        self.progress_bar.setVisible(True)
        self.progress_bar.setRange(0, 0)
        self.progress_bar.setValue(0)
        self.progress_bar.update()
        if message:
            self.status.showMessage(message)

    def stop_progress_indicator(self, message: str | None = None) -> None:
        if self._progress_lock_count > 0:
            self._progress_lock_count -= 1
        if self._progress_lock_count <= 0:
            self._progress_lock_count = 0
            self.progress_bar.setRange(0, 1)
            self.progress_bar.setValue(0)
            self.progress_bar.setVisible(False)
            self.progress_bar.update()
        if message:
            self.status.showMessage(message, 5000)

    def set_busy(self, busy: bool, message: str | None = None) -> None:
        self._full_ui_busy = busy
        if busy:
            self.start_progress_indicator(message)
        else:
            self.stop_progress_indicator(message)
        self._set_main_controls_enabled(not busy)
        if message:
            self.status.showMessage(message)

    def run_worker(
        self,
        fn: Callable[..., Any],
        on_result: Callable[[Any], None],
        busy_message: str,
        *args: Any,
        lock_ui: bool = True,
        done_message: str = "Έτοιμο",
        **kwargs: Any,
    ) -> None:
        if self._is_shutting_down:
            return

        if lock_ui:
            self.set_busy(True, busy_message)
        elif busy_message:
            self.status.showMessage(busy_message, 5000)

        worker = FunctionWorker(fn, *args, **kwargs)
        worker.signals.progress.connect(self.on_worker_progress)
        worker.signals.result.connect(on_result)

        state = {"closed": False}

        def finalize(success: bool) -> None:
            if state["closed"]:
                return
            state["closed"] = True
            if lock_ui:
                self.set_busy(False, done_message if success else None)
            elif success and done_message:
                self.status.showMessage(done_message, 5000)

        def handle_error(payload: Any) -> None:
            finalize(False)
            message, details = self._normalize_worker_error_payload(payload)
            self.show_error(message, details)

        worker.signals.error.connect(handle_error)
        worker.signals.finished.connect(lambda: finalize(True))

        self.thread_pool.start(worker)

    # --------------------------------------------------------
    # Settings / startup
    # --------------------------------------------------------
    def _apply_settings_to_ui(self) -> None:
        self.base_url_edit.setText(self.settings.base_url)
        self.api_key_edit.setText(self.settings.api_key)
        self.timeout_spin.setValue(self.settings.timeout_seconds)
        self.exe_path_edit.setText(self.settings.ollama_exe_path)
        self.cli_path_edit.setText(self.settings.ollama_cli_path)
        self.auto_open_check.setChecked(self.settings.auto_open_ollama_on_start)
        self.auto_start_service_check.setChecked(self.settings.auto_start_service_on_start)
        self.auto_refresh_models_check.setChecked(self.settings.auto_refresh_models_on_start)
        self.auto_load_models_check.setChecked(self.settings.auto_load_selected_models)
        self.collection_name_edit.setText(self.settings.last_collection_name)
        self._set_combo_to_data(self.model_load_scope_combo, self.settings.model_load_scope or MODEL_LOAD_SCOPE_ACTIVE)
        self.chunk_size_spin.setValue(self.settings.chunk_size)
        self.chunk_overlap_spin.setValue(self.settings.chunk_overlap)
        self.embed_batch_spin.setValue(self.settings.embed_batch_size)
        self.top_k_spin.setValue(self.settings.top_k)
        self.temperature_spin.setValue(self.settings.temperature)
        self.max_tokens_spin.setValue(self.settings.max_tokens)
        self._rebuild_model_combos([], preserve=False)
        self._set_combo_to_data(self.chat_mode_combo, "rag" if self.settings.last_chat_mode == CHAT_MODE_RAG else "plain")
        self.on_chat_mode_changed()

    def collect_settings_from_ui(self) -> AppSettings:
        self.settings.base_url = self.base_url_edit.text().strip() or DEFAULT_BASE_URL
        self.settings.api_key = self.api_key_edit.text().strip() or DEFAULT_API_KEY
        self.settings.timeout_seconds = self.timeout_spin.value()
        self.settings.ollama_exe_path = self.exe_path_edit.text().strip()
        self.settings.ollama_cli_path = self.cli_path_edit.text().strip()
        self.settings.auto_open_ollama_on_start = self.auto_open_check.isChecked()
        self.settings.auto_start_service_on_start = self.auto_start_service_check.isChecked()
        self.settings.auto_refresh_models_on_start = self.auto_refresh_models_check.isChecked()
        self.settings.auto_load_selected_models = self.auto_load_models_check.isChecked()
        self.settings.model_load_scope = self.model_load_scope_combo.currentData() or MODEL_LOAD_SCOPE_ACTIVE
        self.settings.last_collection_name = self.collection_name_edit.text().strip() or "knowledge_base"
        self.settings.chunk_size = self.chunk_size_spin.value()
        self.settings.chunk_overlap = self.chunk_overlap_spin.value()
        self.settings.embed_batch_size = self.embed_batch_spin.value()
        self.settings.top_k = self.top_k_spin.value()
        self.settings.temperature = self.temperature_spin.value()
        self.settings.max_tokens = self.max_tokens_spin.value()
        self.settings.last_chat_mode = CHAT_MODE_RAG if self.current_chat_mode_key() == "rag" else CHAT_MODE_PLAIN
        self.settings.preferred_llm = self.llm_combo.currentData() or AUTO_OPTION
        self.settings.preferred_embedding = self.embedding_combo.currentData() or AUTO_OPTION
        self.controller = OllamaController(self.settings)
        self.client = self.controller.client()
        return self.settings

    def save_current_settings(self) -> None:
        settings = self.collect_settings_from_ui()
        save_settings(settings)
        self.log("💾 Οι ρυθμίσεις αποθηκεύτηκαν.")

    def setup_startup_behaviour(self) -> None:
        self.log("🚀 Εκκίνηση εφαρμογής...")
        self.detect_paths(silent=True)
        self.async_refresh_runtime_status()
        self.collect_settings_from_ui()
        save_settings(self.settings)

        def after_ready(_result: Any) -> None:
            self.async_refresh_runtime_status()
            if self.settings.auto_refresh_models_on_start:
                self.refresh_models_clicked()

        self.run_worker(
            self.ensure_ollama_ready_task,
            after_ready,
            "Έλεγχος Ollama κατά την εκκίνηση...",
            lock_ui=False,
            done_message="Η εκκίνηση ολοκληρώθηκε.",
        )

    # --------------------------------------------------------
    # Paths / file selection
    # --------------------------------------------------------
    def select_exe_path(self) -> None:
        path, _ = QFileDialog.getOpenFileName(self, "Επιλογή εκτελέσιμου Ollama")
        if path:
            self.exe_path_edit.setText(path)

    def select_cli_path(self) -> None:
        path, _ = QFileDialog.getOpenFileName(self, "Επιλογή Ollama CLI")
        if path:
            self.cli_path_edit.setText(path)

    def detect_paths(self, silent: bool = False) -> None:
        self.collect_settings_from_ui()
        exe = self.controller.get_effective_exe_path()
        cli = self.controller.get_effective_cli_path()
        if exe:
            self.exe_path_edit.setText(exe)
        if cli:
            self.cli_path_edit.setText(cli)
        if not silent:
            self.log("🔎 Έγινε προσπάθεια εύρεσης path για Ollama / CLI.")

    def select_kb_files(self) -> None:
        paths, _ = QFileDialog.getOpenFileNames(self, "Επιλογή αρχείων για συλλογή γνώσης")
        if not paths:
            return
        for path in paths:
            if path not in self.kb_file_paths:
                self.kb_file_paths.append(path)
        self.refresh_kb_files_list()
        self.log(f"📚 Προστέθηκαν {len(paths)} αρχεία για τη βάση γνώσης.")

    def clear_kb_files(self) -> None:
        self.kb_file_paths.clear()
        self.refresh_kb_files_list()

    def refresh_kb_files_list(self) -> None:
        self.kb_files_list.clear()
        for path in self.kb_file_paths:
            item = QListWidgetItem(path)
            self.kb_files_list.addItem(item)

    def select_chat_attachments(self) -> None:
        paths, _ = QFileDialog.getOpenFileNames(self, "Επιλογή επισυναπτόμενων αρχείων")
        if not paths:
            return
        for path in paths:
            if path not in self.chat_attachment_paths:
                self.chat_attachment_paths.append(path)
        self.refresh_chat_files_list()
        self.log(f"📎 Προστέθηκαν {len(paths)} επισυναπτόμενα αρχεία για την ερώτηση.")

    def clear_chat_attachments(self) -> None:
        self.chat_attachment_paths.clear()
        self.refresh_chat_files_list()

    def ensure_models_cache_available(self, progress_callback: Callable[[str], None]) -> list[AvailableModel]:
        if self.models_cache:
            return self.models_cache
        progress_callback("🔄 Δεν υπάρχει cache μοντέλων. Γίνεται ανανέωση...")
        self.models_cache = self.client.list_models(progress=progress_callback)
        return self.models_cache

    def ensure_server_available(self, progress_callback: Callable[[str], None]) -> None:
        if self.client.is_server_running(timeout=STATUS_CHECK_TIMEOUT):
            return
        if self.controller.ensure_server_running(progress=progress_callback, timeout_seconds=20):
            self.client = self.controller.client()
            return
        raise OllamaError(
            "Το Ollama service δεν απαντά. Ενεργοποίησε την αυτόματη εκκίνηση service ή πάτησε 'Εκκίνηση Ollama Service'."
        )

    def refresh_chat_files_list(self) -> None:
        self.chat_files_list.clear()
        for path in self.chat_attachment_paths:
            self.chat_files_list.addItem(QListWidgetItem(path))

    # --------------------------------------------------------
    # Status / models
    # --------------------------------------------------------
    def status_snapshot_task(
        self,
        settings_snapshot: AppSettings,
        progress_callback: Callable[[str], None] | None = None,
    ) -> dict[str, bool]:
        del progress_callback
        controller = OllamaController(settings_snapshot)
        client = controller.client()
        return {
            "app_running": controller.is_ollama_app_running(),
            "server_running": client.is_server_running(timeout=STATUS_CHECK_TIMEOUT),
        }

    def on_status_snapshot(self, payload: dict[str, bool], log_message: bool = False) -> None:
        self.status_check_in_progress = False
        self.app_status_label.setText("🟢 Ανοιχτό" if payload.get("app_running") else "🔴 Κλειστό")
        self.server_status_label.setText("🟢 Ενεργό" if payload.get("server_running") else "🔴 Ανενεργό")
        if log_message:
            self.log("📡 Έγινε έλεγχος κατάστασης Ollama.")

    def async_refresh_runtime_status(self, log_message: bool = False) -> None:
        if self._is_shutting_down or self.status_check_in_progress:
            return
        settings_snapshot = AppSettings(**asdict(self.collect_settings_from_ui()))
        self.status_check_in_progress = True
        worker = FunctionWorker(self.status_snapshot_task, settings_snapshot)
        worker.signals.result.connect(lambda payload, log_message=log_message: self.on_status_snapshot(payload, log_message))
        worker.signals.error.connect(lambda _msg, log_message=log_message: self.on_status_snapshot({"app_running": False, "server_running": False}, log_message))
        worker.signals.finished.connect(lambda: setattr(self, "status_check_in_progress", False))
        self.thread_pool.start(worker)

    def refresh_runtime_status_quiet(self) -> None:
        self.async_refresh_runtime_status(log_message=False)

    def refresh_runtime_status(self) -> None:
        self.async_refresh_runtime_status(log_message=True)

    def ensure_ollama_ready_task(self, progress_callback: Callable[[str], None]) -> dict[str, bool]:
        return self.controller.ensure_ready(progress=progress_callback)

    def open_ollama_clicked(self) -> None:
        self.collect_settings_from_ui()
        path = self.controller.open_ollama_app()
        self.log(f"🚀 Εκκίνηση Ollama από: {path}")
        QTimer.singleShot(3000, self.async_refresh_runtime_status)

    def start_service_clicked(self) -> None:
        self.collect_settings_from_ui()
        cmd = self.controller.start_service()
        self.log(f"🌐 Εκκίνηση service με: {cmd}")
        QTimer.singleShot(4000, self.async_refresh_runtime_status)

    def refresh_models_clicked(self) -> None:
        self.collect_settings_from_ui()
        self.run_models_worker(
            self.refresh_models_task,
            self.on_models_refreshed,
            "Ανανέωση μοντέλων από το Ollama...",
            done_message="Η ανανέωση μοντέλων ολοκληρώθηκε.",
        )

    def refresh_models_task(self, progress_callback: Callable[[str], None]) -> list[AvailableModel]:
        progress_callback("🔄 Επικοινωνία με το Ollama...")
        self.ensure_server_available(progress_callback)
        models = self.client.list_models(progress=progress_callback)
        progress_callback(f"✅ Βρέθηκαν {len(models)} μοντέλα συνολικά.")
        return models

    def on_models_refreshed(self, models: list[AvailableModel]) -> None:
        self.models_cache = models
        self._rebuild_model_combos(models, preserve=True)
        llm_count = sum(1 for model in models if model.type == "llm")
        embedding_count = sum(1 for model in models if model.type == "embedding")
        self.models_status_label.setText(
            f"LLM: {llm_count} | Embedding: {embedding_count} | Σύνολο: {len(models)}"
        )
        self.log(f"🤖 LLM models: {llm_count} | Embedding models: {embedding_count}")
        self.update_models_list_display()
        self.on_context_for_model_loading_changed()
        self.async_refresh_runtime_status()

    def _rebuild_model_combos(self, models: list[AvailableModel], preserve: bool) -> None:
        previous_llm = self.llm_combo.currentData() if preserve else self.settings.preferred_llm
        previous_embedding = self.embedding_combo.currentData() if preserve else self.settings.preferred_embedding

        self._synchronizing_model_combos = True
        self.llm_combo.clear()
        self.embedding_combo.clear()
        self.llm_combo.addItem(AUTO_OPTION, AUTO_OPTION)
        self.embedding_combo.addItem(AUTO_OPTION, AUTO_OPTION)

        llm_models = [model for model in models if model.type == "llm"]
        embedding_models = [model for model in models if model.type == "embedding"]

        for model in llm_models:
            self.llm_combo.addItem(model.label, model.identifier)
        for model in embedding_models:
            self.embedding_combo.addItem(model.label, model.identifier)

        self._set_combo_to_data(self.llm_combo, previous_llm or AUTO_OPTION)
        self._set_combo_to_data(self.embedding_combo, previous_embedding or AUTO_OPTION)
        self._synchronizing_model_combos = False
        self.update_models_list_display()
        self.on_context_for_model_loading_changed()

    def _set_combo_to_data(self, combo: QComboBox, wanted_data: str) -> None:
        for index in range(combo.count()):
            if combo.itemData(index) == wanted_data:
                combo.setCurrentIndex(index)
                return
        combo.setCurrentIndex(0)

    def current_chat_mode_key(self) -> str:
        return str(self.chat_mode_combo.currentData() or "rag")

    def current_model_load_scope_key(self) -> str:
        if not hasattr(self, "model_load_scope_combo"):
            return MODEL_LOAD_SCOPE_ACTIVE
        return str(self.model_load_scope_combo.currentData() or MODEL_LOAD_SCOPE_ACTIVE)

    def infer_active_model_type(self) -> str:
        current_tab = self.tabs.currentIndex() if hasattr(self, "tabs") else 0
        if current_tab == 1:
            return "embedding"
        return "llm"

    def effective_model_load_scope(self) -> str:
        scope = self.current_model_load_scope_key()
        if scope != MODEL_LOAD_SCOPE_ACTIVE:
            return scope
        return MODEL_LOAD_SCOPE_EMBEDDING if self.infer_active_model_type() == "embedding" else MODEL_LOAD_SCOPE_LLM

    def on_context_for_model_loading_changed(self, *_args: Any) -> None:
        if not hasattr(self, "model_load_scope_combo"):
            return
        active_type = self.infer_active_model_type()
        active_label = "Embedding" if active_type == "embedding" else "LLM"
        tooltip = (
            f"Το ενεργό context αυτή τη στιγμή αντιστοιχεί σε {active_label}. "
            "Στη Βάση Γνώσης φορτώνεται κυρίως embedding model, ενώ στη Συνομιλία/Chat κυρίως LLM."
        )
        self.model_load_scope_combo.setToolTip(tooltip)
        if self.current_model_load_scope_key() == MODEL_LOAD_SCOPE_ACTIVE:
            self.load_models_btn.setText(f"📥 Φόρτωση ενεργού μοντέλου ({active_label})")
        elif self.current_model_load_scope_key() == MODEL_LOAD_SCOPE_LLM:
            self.load_models_btn.setText("📥 Φόρτωση μόνο LLM")
        elif self.current_model_load_scope_key() == MODEL_LOAD_SCOPE_EMBEDDING:
            self.load_models_btn.setText("📥 Φόρτωση μόνο Embedding")
        else:
            self.load_models_btn.setText("📥 Φόρτωση LLM + Embedding")
        self.on_model_selection_changed()

    def set_models_area_busy(self, busy: bool, message: str | None = None) -> None:
        self.models_operation_in_progress = busy
        self._set_models_controls_enabled(not busy)
        if message:
            self.status.showMessage(message)

    def run_models_worker(
        self,
        fn: Callable[..., Any],
        on_result: Callable[[Any], None],
        busy_message: str,
        *args: Any,
        done_message: str = "Τα μοντέλα είναι έτοιμα.",
        **kwargs: Any,
    ) -> None:
        if self._is_shutting_down:
            return

        if self.models_operation_in_progress:
            self.status.showMessage("Υπάρχει ήδη ενεργή εργασία μοντέλων...", 3000)
            return
        self.set_models_area_busy(True, busy_message)
        self.start_progress_indicator(busy_message)

        worker = FunctionWorker(fn, *args, **kwargs)
        worker.signals.progress.connect(self.on_worker_progress)
        worker.signals.result.connect(on_result)

        state = {"closed": False}

        def finalize(success: bool) -> None:
            if state["closed"]:
                return
            state["closed"] = True
            self.set_models_area_busy(False, done_message if success else None)
            self.stop_progress_indicator(done_message if success else None)

        def handle_error(payload: Any) -> None:
            finalize(False)
            message, details = self._normalize_worker_error_payload(payload)
            self.show_error(message, details)

        worker.signals.error.connect(handle_error)
        worker.signals.finished.connect(lambda: finalize(True))
        self.thread_pool.start(worker)

    def on_chat_mode_changed(self, *_args: Any) -> None:
        rag_mode = self.current_chat_mode_key() == "rag"
        self.chat_collection_combo.setEnabled(rag_mode)
        self.top_k_spin.setEnabled(rag_mode)
        self.embedding_combo.setEnabled(rag_mode)
        self.embedding_combo.setToolTip("Χρησιμοποιείται μόνο όταν η συνομιλία είναι σε λειτουργία RAG.")
        if rag_mode:
            self.chat_mode_hint_label.setText(
                "RAG mode: θα χρησιμοποιηθεί η επιλεγμένη συλλογή γνώσης, retrieval Top-K και embedding model. "
                "Η λίστα μοντέλων προσαρμόζεται ώστε να βλέπεις και embeddings."
            )
            self.sources_title_label.setText("Πηγές RAG / συνημμένα")
        else:
            self.chat_mode_hint_label.setText(
                "Chat χωρίς RAG: η ερώτηση στέλνεται απευθείας στο επιλεγμένο LLM. "
                "Δεν απαιτείται συλλογή γνώσης ούτε embedding model και η λίστα μοντέλων φιλτράρεται σε chat/LLM."
            )
            self.sources_title_label.setText("Πηγές / συνημμένα")
        self.update_models_list_display()
        self.on_context_for_model_loading_changed()

    def on_model_selection_changed(self, *_args: Any) -> None:
        llm = self.llm_combo.currentText() or "Αυτόματη επιλογή"
        embedding = self.embedding_combo.currentText() or "Αυτόματη επιλογή"
        scope = self.effective_model_load_scope()
        if scope == MODEL_LOAD_SCOPE_LLM:
            load_hint = "Με το τρέχον context θα φορτώνεται μόνο το LLM."
        elif scope == MODEL_LOAD_SCOPE_EMBEDDING:
            load_hint = "Με το τρέχον context θα φορτώνεται μόνο το embedding model."
        else:
            load_hint = "Έχει επιλεγεί ρητά φόρτωση και των δύο μοντέλων."

        if self.current_chat_mode_key() == "rag":
            self.selected_models_summary.setText(
                f"Ενεργό mode: {CHAT_MODE_RAG}. LLM: {llm}. Embedding: {embedding}. {load_hint}"
            )
        else:
            self.selected_models_summary.setText(
                f"Ενεργό mode: {CHAT_MODE_PLAIN}. LLM: {llm}. Δεν απαιτείται embedding model για το chat. {load_hint}"
            )

    def update_models_list_display(self, *_args: Any) -> None:
        if not hasattr(self, "models_list"):
            return
        filter_mode = self.models_filter_combo.currentData() if hasattr(self, "models_filter_combo") else MODEL_LIST_FILTER_MODE
        if filter_mode == MODEL_LIST_FILTER_MODE:
            filter_mode = MODEL_LIST_FILTER_ALL if self.current_chat_mode_key() == "rag" else MODEL_LIST_FILTER_LLM

        query = self.model_search_edit.text().strip().lower() if hasattr(self, "model_search_edit") else ""
        self.models_list.clear()
        if not self.models_cache:
            self.models_list.addItem("Δεν υπάρχουν ακόμη μοντέλα στην cache. Πάτησε 'Ανανέωση μοντέλων'.")
            return

        filtered_models: list[AvailableModel] = []
        for model in self.models_cache:
            if filter_mode == MODEL_LIST_FILTER_LLM and model.type != "llm":
                continue
            if filter_mode == MODEL_LIST_FILTER_EMBEDDING and model.type != "embedding":
                continue
            searchable = " ".join(
                str(part or "").lower()
                for part in (model.identifier, model.display_name, model.publisher, model.arch, model.type)
            )
            if query and query not in searchable:
                continue
            filtered_models.append(model)

        if not filtered_models:
            self.models_list.addItem("Δεν βρέθηκαν μοντέλα για το τρέχον φίλτρο ή την αναζήτηση.")
            return

        for model in filtered_models:
            state = "φορτωμένο" if model.is_loaded else "διαθέσιμο"
            title = truncate_ui_name(model.display_name or model.identifier, 58)
            publisher = f" • {truncate_ui_name(model.publisher, 18)}" if model.publisher else ""
            self.models_list.addItem(f"{title}  •  {model.type.upper()}{publisher}  •  {state}")

    def resolve_models_preview(self) -> None:
        try:
            llm_model, embedding_model = self.resolve_models_for_current_state()
            self.log(f"🤖 Auto-detect LLM: {llm_model}")
            self.log(f"🧠 Auto-detect Embedding: {embedding_model}")
            self.show_info(
                f"Auto-detect αποτέλεσμα:\n\nLLM: {llm_model}\nEmbedding: {embedding_model}"
            )
        except Exception as exc:
            self.show_error(str(exc))

    def resolve_models_for_current_state(
        self,
        collection_name: str | None = None,
        requested_llm: str | None = None,
        requested_embedding: str | None = None,
        require_embedding: bool = True,
    ) -> tuple[str, str]:
        if not self.models_cache:
            raise OllamaError("Δεν υπάρχουν μοντέλα στην cache. Πάτησε πρώτα 'Ανανέωση μοντέλων'.")

        preferred_embedding = None
        if collection_name:
            try:
                metadata = self.engine.get_collection_metadata(collection_name)
                preferred_embedding = metadata.get("embedding_model")
            except Exception:
                preferred_embedding = None

        llm_model = resolve_model_identifier(
            models=self.models_cache,
            model_type="llm",
            requested_identifier=requested_llm or self.llm_combo.currentData(),
            preferred_identifier=self.settings.preferred_llm,
        )
        embedding_model = ""
        if require_embedding:
            embedding_model = resolve_model_identifier(
                models=self.models_cache,
                model_type="embedding",
                requested_identifier=requested_embedding or self.embedding_combo.currentData(),
                preferred_identifier=preferred_embedding or self.settings.preferred_embedding,
            )
        return llm_model, embedding_model

    def load_selected_models_clicked(self) -> None:
        self.collect_settings_from_ui()
        self.run_models_worker(
            self.load_selected_models_task,
            self.on_models_loaded,
            "Φόρτωση μοντέλου στο Ollama...",
            self.llm_combo.currentData(),
            self.embedding_combo.currentData(),
            self.current_model_load_scope_key(),
            done_message="Η φόρτωση μοντέλου ολοκληρώθηκε.",
        )

    def load_selected_models_task(
        self,
        requested_llm: str | None,
        requested_embedding: str | None,
        requested_scope: str,
        progress_callback: Callable[[str], None],
    ) -> dict[str, str]:
        self.ensure_server_available(progress_callback)
        self.ensure_models_cache_available(progress_callback)
        scope = requested_scope or MODEL_LOAD_SCOPE_ACTIVE
        if scope == MODEL_LOAD_SCOPE_ACTIVE:
            scope = self.effective_model_load_scope()

        llm_model = ""
        embedding_model = ""

        if scope in (MODEL_LOAD_SCOPE_LLM, MODEL_LOAD_SCOPE_BOTH):
            llm_model = resolve_model_identifier(
                models=self.models_cache,
                model_type="llm",
                requested_identifier=requested_llm,
                preferred_identifier=self.settings.preferred_llm,
            )
            progress_callback(f"📥 Load LLM: {llm_model}")
            self.client.load_model(llm_model, "llm")

        if scope in (MODEL_LOAD_SCOPE_EMBEDDING, MODEL_LOAD_SCOPE_BOTH):
            embedding_model = resolve_model_identifier(
                models=self.models_cache,
                model_type="embedding",
                requested_identifier=requested_embedding,
                preferred_identifier=self.settings.preferred_embedding,
            )
            progress_callback(f"📥 Load Embedding: {embedding_model}")
            self.client.load_model(embedding_model, "embedding")

        return {"llm": llm_model, "embedding": embedding_model, "scope": scope}

    def on_models_loaded(self, payload: dict[str, str]) -> None:
        if payload.get('llm'):
            self.log(f"✅ Φορτώθηκε LLM: {payload['llm']}")
        if payload.get('embedding'):
            self.log(f"✅ Φορτώθηκε Embedding: {payload['embedding']}")
        if not payload.get('llm') and not payload.get('embedding'):
            self.log("ℹ️ Δεν ζητήθηκε φόρτωση μοντέλου.")
        QTimer.singleShot(80, self.refresh_models_clicked)

    # --------------------------------------------------------
    # Collections
    # --------------------------------------------------------
    def refresh_collections(self) -> None:
        previous_collection = self.chat_collection_combo.currentText().strip()
        current_list_item = self.collections_list.currentItem()
        previous_list_name = current_list_item.text() if current_list_item else ""

        collections = self.engine.list_collections()
        self.collections_list.clear()
        self.chat_collection_combo.clear()
        self.chat_collection_combo.addItem("")
        for name in collections:
            self.collections_list.addItem(name)
            self.chat_collection_combo.addItem(name)

        if previous_collection:
            index = self.chat_collection_combo.findText(previous_collection)
            if index >= 0:
                self.chat_collection_combo.setCurrentIndex(index)

        if previous_list_name:
            matches = self.collections_list.findItems(previous_list_name, Qt.MatchFlag.MatchExactly)
            if matches:
                self.collections_list.setCurrentItem(matches[0])
            else:
                self.collection_metadata_output.clear()
        else:
            self.collection_metadata_output.clear()

        self.log(f"📚 Συλλογές διαθέσιμες: {len(collections)}")
        self.on_chat_mode_changed()

    def show_selected_collection_metadata(
        self,
        current: QListWidgetItem | None = None,
        previous: QListWidgetItem | None = None,
    ) -> None:
        del previous
        item = current or self.collections_list.currentItem()
        if not item:
            self.collection_metadata_output.clear()
            return
        name = item.text()
        try:
            metadata = self.engine.get_collection_metadata(name)
            self.collection_metadata_output.setPlainText(
                json.dumps(metadata, ensure_ascii=False, indent=2)
            )
        except Exception as exc:
            self.collection_metadata_output.setPlainText(str(exc))

    def delete_selected_collection(self) -> None:
        item = self.collections_list.currentItem()
        if not item:
            self.show_error("Επίλεξε πρώτα μια συλλογή για διαγραφή.")
            return
        name = item.text()
        reply = QMessageBox.question(
            self,
            APP_NAME,
            f"Θέλεις σίγουρα να διαγράψεις τη συλλογή '{name}';",
        )
        if reply != QMessageBox.StandardButton.Yes:
            return
        self.engine.delete_collection(name)
        self.refresh_collections()
        self.log(f"🗑️ Διαγράφηκε η συλλογή: {name}")

    def build_collection_clicked(self) -> None:
        collection_name = self.collection_name_edit.text().strip()
        if not collection_name:
            self.show_error("Δώσε όνομα συλλογής.")
            return
        if not self.kb_file_paths:
            self.show_error("Επίλεξε πρώτα αρχεία για τη συλλογή.")
            return
        self.collect_settings_from_ui()
        request = {
            "collection_name": collection_name,
            "file_paths": list(self.kb_file_paths),
            "requested_llm": self.llm_combo.currentData(),
            "requested_embedding": self.embedding_combo.currentData(),
            "chunk_size": self.chunk_size_spin.value(),
            "chunk_overlap": self.chunk_overlap_spin.value(),
            "batch_size": self.embed_batch_spin.value(),
            "auto_load": self.settings.auto_load_selected_models,
        }
        self.run_worker(
            self.build_collection_task,
            self.on_collection_built,
            "Δημιουργία συλλογής γνώσης...",
            request,
        )

    def build_collection_task(
        self,
        request: dict[str, Any],
        progress_callback: Callable[[str], None],
    ) -> dict[str, Any]:
        self.ensure_server_available(progress_callback)
        self.ensure_models_cache_available(progress_callback)
        collection_name = request["collection_name"]
        _llm_model, embedding_model = self.resolve_models_for_current_state(
            collection_name=collection_name,
            requested_llm=request.get("requested_llm"),
            requested_embedding=request.get("requested_embedding"),
        )

        if request.get("auto_load"):
            progress_callback(f"📥 Auto-load embedding model: {embedding_model}")
            self.client.load_model(embedding_model, "embedding")

        metadata = self.engine.build_collection(
            collection_name=collection_name,
            file_paths=request["file_paths"],
            client=self.client,
            embedding_model=embedding_model,
            chunk_size=int(request["chunk_size"]),
            chunk_overlap=int(request["chunk_overlap"]),
            batch_size=int(request["batch_size"]),
            progress=progress_callback,
        )
        return metadata

    def on_collection_built(self, metadata: dict[str, Any]) -> None:
        self.log(f"✅ Η συλλογή '{metadata['collection_name']}' είναι έτοιμη.")
        self.refresh_collections()
        self.collection_name_edit.setText(metadata["collection_name"])
        self.async_refresh_runtime_status()

    # --------------------------------------------------------
    # Chat / question answering
    # --------------------------------------------------------
    def ask_question_clicked(self) -> None:
        question = self.question_input.toPlainText().strip()
        if not question:
            self.show_error("Γράψε πρώτα μια ερώτηση.")
            return

        chat_mode = self.current_chat_mode_key()
        collection_name = self.chat_collection_combo.currentText().strip()
        if chat_mode == "rag" and not collection_name:
            self.show_error("Στη λειτουργία RAG πρέπει να επιλέξεις συλλογή γνώσης.")
            return

        self.collect_settings_from_ui()
        request = {
            "question": question,
            "chat_mode": chat_mode,
            "collection_name": collection_name,
            "use_collection": chat_mode == "rag",
            "task_mode": self.task_mode_combo.currentText().strip(),
            "attachment_paths": list(self.chat_attachment_paths),
            "requested_llm": self.llm_combo.currentData(),
            "requested_embedding": self.embedding_combo.currentData(),
            "auto_load": self.settings.auto_load_selected_models,
            "top_k": self.top_k_spin.value(),
            "temperature": self.temperature_spin.value(),
            "max_tokens": self.max_tokens_spin.value(),
        }
        self.run_worker(
            self.ask_question_task,
            self.on_question_answered,
            "Υποβολή ερώτησης στο Ollama...",
            request,
        )

    def ask_question_task(
        self,
        request: dict[str, Any],
        progress_callback: Callable[[str], None],
    ) -> dict[str, Any]:
        self.ensure_server_available(progress_callback)
        self.ensure_models_cache_available(progress_callback)

        question = request["question"]
        chat_mode = str(request.get("chat_mode") or "rag").strip().lower()
        collection_name = (request.get("collection_name") or "").strip()
        use_collection = chat_mode == "rag" and bool(request.get("use_collection")) and bool(collection_name)
        task_mode = (request.get("task_mode") or "Αυτόματο").strip()
        attached_docs: list[LoadedDocument] = []

        progress_callback("📎 Ανάγνωση επισυναπτόμενων αρχείων...")
        for path in request.get("attachment_paths", []):
            if not Path(path).exists():
                raise DocumentLoaderError(f"Δεν βρέθηκε το συνημμένο αρχείο: {path}")
            attached_docs.append(extract_text_from_path(path))

        if task_mode == "Αυτόματο" and attached_docs:
            if any(Path(doc.filename).suffix.lower() in CODE_FILE_EXTENSIONS for doc in attached_docs):
                task_mode = "Έλεγχος κώδικα"
            else:
                task_mode = "Εξήγηση αρχείων"

        llm_model = resolve_model_identifier(
            models=self.models_cache,
            model_type="llm",
            requested_identifier=request.get("requested_llm"),
            preferred_identifier=self.settings.preferred_llm,
        )

        embedding_model = ""
        if use_collection:
            llm_model, embedding_model = self.resolve_models_for_current_state(
                collection_name=collection_name,
                requested_llm=request.get("requested_llm"),
                requested_embedding=request.get("requested_embedding"),
            )

        if request.get("auto_load"):
            progress_callback(f"📥 Auto-load LLM: {llm_model}")
            self.client.load_model(llm_model, "llm")
            if use_collection and embedding_model:
                progress_callback("🧠 Το embedding model θα χρησιμοποιηθεί on-demand μόνο αν χρειαστεί για retrieval.")

        rag_hits: list[SearchResult] = []
        if use_collection:
            progress_callback(f"🔎 Αναζήτηση σε συλλογή: {collection_name}")
            rag_hits = self.engine.search(
                collection_name=collection_name,
                query=question,
                client=self.client,
                embedding_model=embedding_model,
                top_k=int(request.get("top_k", DEFAULT_TOP_K)),
            )

        progress_callback("💬 Δημιουργία prompt και κλήση chat completion...")
        messages = build_chat_messages(
            question=question,
            rag_hits=rag_hits,
            attached_docs=attached_docs,
            task_mode=task_mode,
        )
        answer = self.client.chat(
            model=llm_model,
            messages=messages,
            temperature=float(request.get("temperature", DEFAULT_TEMPERATURE)),
            max_tokens=int(request.get("max_tokens", DEFAULT_MAX_TOKENS)),
        )

        return {
            "answer": answer,
            "sources": [asdict(hit) for hit in rag_hits],
            "llm_model": llm_model,
            "embedding_model": embedding_model,
            "task_mode": task_mode,
            "chat_mode": chat_mode,
            "attachments": [asdict(doc) for doc in attached_docs],
        }

    def on_question_answered(self, payload: dict[str, Any]) -> None:
        answer = payload["answer"]
        self.latest_answer = answer
        self.answer_output.setMarkdown(answer)
        self.answer_output.moveCursor(QTextCursor.MoveOperation.Start)

        sources = payload.get("sources", [])
        attachments = payload.get("attachments", [])
        if sources:
            rendered_sources = []
            for idx, src in enumerate(sources, start=1):
                rendered_sources.append(
                    f"[ΠΗΓΗ {idx}] αρχείο={src['filename']} | chunk={src['chunk_index']} | score={src['score']:.4f}\n{src['text']}"
                )
            self.sources_output.setPlainText("\n\n".join(rendered_sources))
        elif attachments:
            attachment_lines = [f"- {item.get('filename', 'άγνωστο αρχείο')}" for item in attachments]
            self.sources_output.setPlainText(
                "Δεν έγινε retrieval από συλλογή RAG. Η απάντηση βασίστηκε στα παρακάτω επισυναπτόμενα αρχεία:\n"
                + "\n".join(attachment_lines)
            )
        else:
            self.sources_output.setPlainText("Η απάντηση δόθηκε χωρίς RAG πηγές και χωρίς επισυναπτόμενα αρχεία.")

        emb = payload['embedding_model'] or '—'
        mode_label = CHAT_MODE_RAG if payload.get('chat_mode') == 'rag' else CHAT_MODE_PLAIN
        self.log(
            f"✅ Απάντηση ολοκληρώθηκε | συνομιλία={mode_label} | task={payload['task_mode']} | LLM={payload['llm_model']} | Embedding={emb}"
        )

    def clear_answer_output(self) -> None:
        self.answer_output.clear()
        self.sources_output.clear()
        self.latest_answer = ""

    def copy_answer_to_clipboard(self) -> None:
        text = self.answer_output.toPlainText().strip()
        if not text:
            self.show_error("Δεν υπάρχει απάντηση για αντιγραφή.")
            return
        QApplication.clipboard().setText(text)
        self.status.showMessage("Η απάντηση αντιγράφηκε στο clipboard.", 3000)
        self.log("📋 Η απάντηση αντιγράφηκε στο clipboard.")

    def save_answer_text(self) -> None:
        text = self.answer_output.toPlainText().strip()
        if not text:
            self.show_error("Δεν υπάρχει απάντηση για αποθήκευση.")
            return
        path, _ = QFileDialog.getSaveFileName(self, "Αποθήκευση απάντησης", "answer.txt", "Text Files (*.txt)")
        if not path:
            return
        Path(path).write_text(text, encoding="utf-8")
        self.log(f"💾 Αποθηκεύτηκε η απάντηση στο: {path}")

    def save_corrected_code(self) -> None:
        answer = self.answer_output.toPlainText().strip()
        if not answer:
            self.show_error("Δεν υπάρχει απάντηση για εξαγωγή code.")
            return
        language, code = extract_first_code_block(answer)
        if not code:
            self.show_error("Δεν βρέθηκε fenced code block μέσα στην απάντηση.")
            return

        default_name = "corrected_code"
        if self.chat_attachment_paths:
            default_name = Path(self.chat_attachment_paths[0]).name
        extension = f"*.{language}" if language else "*.*"
        path, _ = QFileDialog.getSaveFileName(
            self,
            "Αποθήκευση corrected code",
            default_name,
            f"Code Files ({extension});;All Files (*.*)",
        )
        if not path:
            return
        Path(path).write_text(code, encoding="utf-8")
        self.log(f"🧩 Αποθηκεύτηκε corrected code στο: {path}")


    def closeEvent(self, event: QCloseEvent) -> None:
        self._is_shutting_down = True

        for timer_name in ("log_flush_timer", "status_timer", "progress_watchdog_timer"):
            timer = getattr(self, timer_name, None)
            if timer is None:
                continue
            try:
                timer.stop()
            except RuntimeError:
                pass

        try:
            self.save_current_settings()
        except Exception:
            pass

        try:
            self.thread_pool.clear()
        except Exception:
            pass

        try:
            self.thread_pool.waitForDone(1200)
        except TypeError:
            try:
                self.thread_pool.waitForDone()
            except Exception:
                pass
        except Exception:
            pass

        try:
            super().closeEvent(event)
        except Exception:
            event.accept()


# ============================================================
# Εκκίνηση εφαρμογής
# ============================================================
def main() -> int:
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    current_font = app.font()
    safe_family = current_font.family() or "Segoe UI"
    safe_font = QFont(safe_family)
    safe_point_size = current_font.pointSizeF()
    safe_font.setPointSizeF(max(11.0, safe_point_size if safe_point_size and safe_point_size > 0 else 11.0))
    app.setFont(safe_font)
    app.setApplicationName(APP_NAME)
    app.setOrganizationName("OpenAI")
    window = MainWindow()
    window.showMaximized()
    return app.exec()


if __name__ == "__main__":
    raise SystemExit(main())
